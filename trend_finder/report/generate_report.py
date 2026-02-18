"""
Generate TrendPilot Comprehensive Progress Report as a Word document.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# ── Helper functions ──
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

def add_code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(40, 40, 40)
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.left_indent = Cm(1)
    return p

# ═══════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('TrendPilot')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('AI-Powered LinkedIn Content Strategy Tool')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()

line = doc.add_paragraph()
line.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = line.add_run('_' * 50)
run.font.color.rgb = RGBColor(150, 150, 150)

doc.add_paragraph()

report_title = doc.add_paragraph()
report_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = report_title.add_run('Comprehensive Progress Report')
run.bold = True
run.font.size = Pt(20)

doc.add_paragraph()

for line_text in [
    'Author: Biraj Mishra',
    'Date: February 17, 2026',
    'Course: MS Capstone Project',
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line_text)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════
add_heading_styled('Table of Contents', level=1)
toc_items = [
    '1.  Executive Summary',
    '2.  Project Overview and Architecture',
    '3.  Module 1: Trending Topic Identification',
    '    3.1  Objective',
    '    3.2  Iteration 1: BuzzSumo + News API',
    '    3.3  Iteration 2: Google Trends + spaCy NLP',
    '    3.4  Iteration 3: Google Trends + GPT-4o Prompt Engineering',
    '    3.5  Iteration 4: Interactive Loop + Version Safety + Query Fix',
    '    3.6  Current System Architecture',
    '    3.7  Key Learnings',
    '4.  Module 2: Post Generation',
    '    4.1  Overview and Approach',
    '    4.2  System Prompt Design',
    '    4.3  User Prompt Builder',
    '    4.4  Multi-Variant Generation',
    '    4.5  Sample Output',
    '5.  Module 3: Engagement Prediction',
    '    5.1  Abstract and Objectives',
    '    5.2  Data Collection and Preprocessing',
    '    5.3  Feature Engineering (85 Features)',
    '    5.4  Model Training and Selection',
    '    5.5  Results and Feature Importance',
    '    5.6  Validation and Robustness',
    '    5.7  Key Findings',
    '6.  Module 4: Visual Generation & Publishing Support',
    '    6.1  Architecture Overview',
    '    6.2  LLM Classification Layer',
    '    6.3  Diagram Generation (Mermaid)',
    '    6.4  Image Generation (Stable Diffusion)',
    '    6.5  Service Orchestration',
    '7.  Integration Architecture',
    '8.  Current Status Summary',
    '9.  Pending Work and Next Steps',
    '10. Conclusion',
    '11. References',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if item.startswith('    '):
        p.paragraph_format.left_indent = Cm(1.5)
        for run in p.runs:
            run.font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════
add_heading_styled('1. Executive Summary', level=1)

add_body(
    'TrendPilot is a modular, agent-based system designed to automate end-to-end professional '
    'content creation for LinkedIn. The system addresses a critical gap in the content creation '
    'workflow: professionals spend significant time manually discovering trends, crafting posts, '
    'predicting what will resonate, and creating supporting visuals. TrendPilot automates this '
    'entire pipeline through four integrated modules.'
)

add_body(
    'This report documents the comprehensive progress across all four modules of the TrendPilot '
    'system, covering iterative development decisions, technical architecture, experimental results, '
    'and remaining work items. The project has achieved substantial milestones across all modules, '
    'with core algorithmic development complete and integration/testing phases pending.'
)

add_heading_styled('Key Achievements', level=2)
achievements = [
    ('Module 1 (Trend Identification): ', 'Four iterations completed, evolving from BuzzSumo + News API '
     'to a GPT-4o + Google Trends pipeline with interactive topic selection and hallucination guards.'),
    ('Module 2 (Post Generation): ', 'LLM-based post generation pipeline implemented with master system '
     'prompt, parameterized user prompts, and multi-variant generation across hook styles.'),
    ('Module 3 (Engagement Prediction): ', 'Production-ready ML models achieving R\u00b2 = 0.5903 for reactions '
     '(Random Forest) and R\u00b2 = 0.5280 for comments (LightGBM), trained on 31,996 posts from 69 influencers '
     'with 85 engineered features.'),
    ('Module 4 (Visual Generation): ', 'Unified visual service integrating LLM classification, Mermaid diagram '
     'rendering, and Stable Diffusion image generation via a CLI-based orchestration layer.'),
]
for bold_part, rest in achievements:
    add_bullet(rest, bold_prefix=bold_part)

add_heading_styled('Pending Work', level=2)
pending = [
    'Unit testing across all modules',
    'Integration testing between modules',
    'Streamlit application development for unified user interface',
    'End-to-end pipeline orchestration',
]
for item in pending:
    add_bullet(item)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. PROJECT OVERVIEW AND ARCHITECTURE
# ═══════════════════════════════════════════════════════════════
add_heading_styled('2. Project Overview and Architecture', level=1)

add_body(
    'TrendPilot is architected as a modular pipeline where each module operates as an independent '
    'agent with well-defined inputs and outputs. This design enables independent development, testing, '
    'and enhancement of each component without impacting downstream modules.'
)

add_heading_styled('2.1 System Pipeline', level=2)
add_body('The end-to-end TrendPilot pipeline follows this flow:')

pipeline_steps = [
    ('Step 1 \u2013 Identify Trending Topics: ', 'Analyze user\'s LinkedIn bio, extract professional topics, '
     'query Google Trends for real-time search data, rank by trend score.'),
    ('Step 2 \u2013 Generate Multiple Posts: ', 'Take selected trending topic, apply master system prompt with '
     'engagement optimization rules, generate multiple post variants with different hook styles.'),
    ('Step 3 \u2013 Predict Engagement: ', 'Score each generated post variant using ML models trained on 31,996 '
     'LinkedIn posts, predicting reactions and comments.'),
    ('Step 4 \u2013 Generate Visuals: ', 'Classify visual need (diagram vs. image), generate appropriate visual '
     'content using Mermaid or Stable Diffusion.'),
    ('Step 5 \u2013 Publish or Review: ', 'Present ranked posts with predicted engagement scores and generated '
     'visuals for user review and publishing decision.'),
]
for bold_part, rest in pipeline_steps:
    add_bullet(rest, bold_prefix=bold_part)

add_heading_styled('2.2 Technology Stack', level=2)
add_table(
    ['Component', 'Technology', 'Purpose'],
    [
        ['LLM (Topic Extraction)', 'GPT-4o (OpenAI)', 'Professional topic extraction and post generation'],
        ['Trend Data', 'Google Trends (pytrends)', 'Real-time search interest and related queries'],
        ['NLP Processing', 'spaCy, VADER, textstat', 'NER, sentiment analysis, readability metrics'],
        ['ML Models', 'Random Forest, LightGBM', 'Engagement prediction (reactions and comments)'],
        ['Diagram Generation', 'Mermaid.js', 'Flowcharts, sequence diagrams, architecture diagrams'],
        ['Image Generation', 'Stable Diffusion', 'AI-generated images for post visuals'],
        ['Frontend (Planned)', 'Streamlit', 'Interactive web application'],
        ['Language', 'Python 3.x', 'Core implementation language'],
    ]
)

add_heading_styled('2.3 Data Sources', level=2)
add_table(
    ['Source', 'Type', 'Volume', 'Usage'],
    [
        ['Google Trends API', 'Search interest data', 'Real-time', 'Trend scoring and related queries'],
        ['LinkedIn Influencer Dataset (Kaggle)', 'Social media posts', '31,996 posts / 69 influencers', 'Engagement prediction training'],
        ['NewsAPI.org', 'News headlines', '30,000+ articles', 'Initial trend discovery (Iter 1)'],
        ['GNews API', 'News headlines', 'Supplementary', 'Secondary source for trend validation'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. MODULE 1: TRENDING TOPIC IDENTIFICATION
# ═══════════════════════════════════════════════════════════════
add_heading_styled('3. Module 1: Trending Topic Identification', level=1)

add_heading_styled('3.1 Objective', level=2)
add_body(
    'Build a system that takes a LinkedIn professional\'s profile bio and automatically identifies '
    'timely, specific, trending topics relevant to their expertise \u2013 then recommends a concrete '
    'LinkedIn post angle backed by real-time search data from Google Trends.'
)
add_body(
    'The system must produce post recommendations that are specific enough to be immediately actionable '
    '(e.g., referencing a new product launch or rising search trend), not generic thought-leadership '
    'angles (e.g., "Why AI is the future"). The recommendations must be grounded in verifiable, '
    'real-time data.'
)

add_heading_styled('Iteration Summary', level=3)
add_table(
    ['Iteration', 'Approach', 'Topic Extraction', 'Trend Source', 'Outcome'],
    [
        ['1', 'BuzzSumo + News API', 'Keyword matching', 'News articles', 'Generic, noisy'],
        ['2', 'Google Trends + spaCy', 'Named Entity Recognition', 'related_queries', 'Ambiguous entities'],
        ['3', 'Google Trends + GPT-4o', 'LLM prompt engineering', 'interest_over_time + related_queries', 'Specific, actionable'],
        ['4', 'Iter 3 + Interactive', 'Same as Iter 3', 'Same + merged API calls', 'User control, factual'],
    ]
)

# --- Iteration 1 ---
add_heading_styled('3.2 Iteration 1: BuzzSumo + News API', level=2)
add_body(
    'The initial approach used BuzzSumo API to find popular content and News API to find trending articles. '
    'Keywords were extracted from the user\'s LinkedIn bio using simple keyword matching, then matched '
    'against news headlines and BuzzSumo trending content.'
)

add_heading_styled('Problems Encountered', level=3)
problems_iter1 = [
    ('Generic topics: ', 'Results were dominated by broad news stories (e.g., "AI is transforming business") '
     'with no specific angle for a LinkedIn post.'),
    ('Noisy results: ', 'News API returned articles about politics, sports, entertainment \u2013 not filtered '
     'to the user\'s professional domain.'),
    ('Low relevance: ', 'No mechanism to score how relevant a trending topic was to the user\'s specific skill set.'),
]
for bold_part, rest in problems_iter1:
    add_bullet(rest, bold_prefix=bold_part)

add_body(
    'Decision: Abandon news-based approach. Shift to Google Trends for real-time search interest data '
    'that better reflects what professionals are actually searching for.'
)

# --- Iteration 2 ---
add_heading_styled('3.3 Iteration 2: Google Trends + spaCy NLP', level=2)
add_body('File: trend_finder.py (596 lines)')
add_body(
    'Used spaCy\'s en_core_web_sm NLP model to extract named entities (ORG, PRODUCT, PROPN, etc.) from '
    'the user\'s LinkedIn bio, then queried Google Trends related_queries API for each extracted entity.'
)

add_heading_styled('Key Components', level=3)
components = [
    'Entity extraction using spaCy NER with 11 entity labels (PERSON, ORG, GPE, PRODUCT, EVENT, etc.)',
    'Proper noun extraction for tokens tagged as PROPN',
    'Noun phrase extraction for multi-word technical terms',
    'TrendCache class with 7-day TTL and composite cache keys (keyword|geo|timeframe)',
    'Retry logic with exponential backoff for Google Trends rate limiting (HTTP 429 errors)',
]
for c in components:
    add_bullet(c)

add_heading_styled('Critical Problems', level=3)
add_body(
    'spaCy NER proved fundamentally unsuitable for technical topic extraction from professional bios. '
    'The model lacks domain context and cannot distinguish ambiguous terms:'
)

add_table(
    ['Entity Extracted', 'Expected Context', 'Google Trends Results', 'Root Cause'],
    [
        ['"Cloud" (PROPN)', 'Cloud computing', 'cloud weather today, cloud types, cumulus cloud', 'No domain awareness'],
        ['"Spring" (PROPN)', 'Spring Boot/Cloud', 'spring 2026, spring fever, hot spring', 'No tech context'],
        ['"Kong" (ORG)', 'Kong API Gateway', 'king kong, hong kong china, donkey kong', 'Pop culture dominance'],
    ]
)

add_body(
    'Additionally, aggressive rate limiting (HTTP 429) from Google Trends API caused data loss. '
    'Decision: Replace spaCy NER entirely with LLM-based prompt engineering.'
)

# --- Iteration 3 ---
add_heading_styled('3.4 Iteration 3: Google Trends + GPT-4o Prompt Engineering', level=2)
add_body('File: trend_identification_v2.py (initial version)')
add_body(
    'Replaced spaCy NER entirely with a GPT-4o prompt that understands professional context. Added '
    'interest_over_time API for trend scoring alongside related_queries (top + rising) for search context. '
    'The LLM both extracts topics AND recommends a post angle.'
)

add_heading_styled('Topic Extraction Prompt', level=3)
add_code_block(
    'You are analyzing a LinkedIn professional bio.\n\n'
    'Extract 8-12 specific, post-worthy technical topics\n'
    'suitable for LinkedIn.\n'
    'Only include:\n'
    '- Technologies, Platforms, Tools, Frameworks,\n'
    '  Engineering practices\n\n'
    'Exclude:\n'
    '- Job titles\n'
    '- Generic words (e.g., cloud, data, software)\n'
    '- Soft skills\n'
    '- Single generic nouns\n\n'
    'Return the result as a comma-separated list.'
)

add_body(
    'This prompt explicitly excludes generic words like "cloud", "data", "software" \u2013 solving the '
    'Iteration 2 ambiguity problem. The LLM understands that "Spring Cloud" is a technology while '
    '"Spring" alone is ambiguous.'
)

add_heading_styled('Remaining Issues', level=3)
add_bullet('Blank related queries due to double build_payload() calls causing rate limiting', bold_prefix='API Rate Limiting: ')
add_bullet(
    'When recommending a post about OpenShift, the LLM produced "OpenShift 4.12" \u2013 but the '
    'current version was 4.21. The version was fabricated from stale training data.',
    bold_prefix='LLM Hallucination: '
)

# --- Iteration 4 ---
add_heading_styled('3.5 Iteration 4: Interactive Loop + Version Safety + Query Fix', level=2)
add_body('File: trend_identification_v2.py (current version, 272 lines)')

fixes = [
    ('Fix 1 \u2013 Merged API Calls: ', 'Combined interest_over_time and related_queries into a single '
     'fetch_trend_data() function with one build_payload() call per keyword, eliminating rate-limiting '
     'on the second call. Cache validation updated to re-fetch entries missing query data.'),
    ('Fix 2 \u2013 Version Hallucination Guard: ', 'Added critical prompt rules: "NEVER invent or guess '
     'version numbers. ONLY mention a specific version if it explicitly appears in the search query data. '
     'If no version is mentioned, use phrasing like latest release or new update."'),
    ('Fix 3 \u2013 Interactive Topic Selection: ', 'Users now see top 5 topics with trend scores and can '
     'interactively choose which topic to explore. After each recommendation, the menu loops back for '
     'multi-topic exploration without restarting.'),
    ('Fix 4 \u2013 Error Visibility: ', 'Replaced silent exception swallowing (except Exception: pass) with '
     'explicit error printing for debugging.'),
    ('Fix 5 \u2013 JSON Serialization: ', 'Added explicit type conversion for numpy int64 values to prevent '
     'serialization errors when saving to JSON cache.'),
]
for bold_part, rest in fixes:
    add_bullet(rest, bold_prefix=bold_part)

# --- Architecture ---
add_heading_styled('3.6 Current System Architecture', level=2)
add_body('The current pipeline follows a two-stage LLM architecture with Google Trends data enrichment:')

arch_steps = [
    'LinkedIn Bio (User Input) \u2013 User enters professional bio text',
    'GPT-4o Topic Extraction \u2013 Extract 8\u201312 specific technical topics',
    'Google Trends API (pytrends) \u2013 interest_over_time + related_queries',
    'Rank & Display Top 5 Topics \u2013 Sort by trend score, show search context',
    'User Selection (Interactive) \u2013 Pick 1\u20135, AI auto-pick, or quit',
    'GPT-4o Post Recommendation \u2013 Specific, timely title (no hallucinated versions)',
]
for i, step in enumerate(arch_steps, 1):
    add_bullet(step, bold_prefix=f'Stage {i}: ')

# --- Key Learnings ---
add_heading_styled('3.7 Key Learnings', level=2)
add_table(
    ['Problem', 'Root Cause', 'Solution'],
    [
        ['Generic topics (Iter 1)', 'News APIs return broad content', 'Switched to Google Trends'],
        ['"Cloud" = weather (Iter 2)', 'spaCy NER has no domain awareness', 'GPT-4o prompt excludes generic words'],
        ['"Spring" = season (Iter 2)', 'NLP model lacks tech context', 'LLM extracts "Spring Cloud" as whole term'],
        ['Blank queries (Iter 3)', 'Double build_payload() rate limited', 'Merged into single fetch_trend_data()'],
        ['Hallucinated versions (Iter 3)', 'LLM invents from stale training data', 'Prompt: only cite versions from search data'],
        ['One-shot exit (Iter 3)', 'Program terminated after one run', 'Interactive loop with re-selection'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. MODULE 2: POST GENERATION
# ═══════════════════════════════════════════════════════════════
add_heading_styled('4. Module 2: Post Generation', level=1)

add_heading_styled('4.1 Overview and Approach', level=2)
add_body(
    'The Post Generation module takes trending topics identified by Module 1 and generates high-quality, '
    'engagement-optimized LinkedIn posts using LLM-based content generation. The module is implemented '
    'as a Jupyter notebook prototype (Post_generator.ipynb) using the OpenAI GPT-4o-mini API, with plans '
    'for integration into the Streamlit application as a dedicated agent.'
)
add_body(
    'The approach uses a two-prompt architecture: a master system prompt that defines the content '
    'strategy rules, and a parameterized user prompt that injects topic-specific context, personal '
    'insights, and hook style preferences.'
)

add_heading_styled('4.2 System Prompt Design', level=2)
add_body(
    'The master system prompt encodes LinkedIn content best practices derived from engagement research '
    'and platform algorithm understanding:'
)

add_code_block(
    'You are an expert LinkedIn content strategist and copywriter.\n\n'
    'Your goal is to generate authentic, high-engagement LinkedIn posts\n'
    'that maximize reach and interaction while avoiding promotional suppression.\n\n'
    'Optimize for:\n'
    '- Authentic storytelling over promotion\n'
    '- Strong first-sentence hooks\n'
    '- Specific details (numbers, moments, experiences)\n'
    '- Personal insights or transformations\n\n'
    'Hard constraints:\n'
    '- Target length: 100-200 words\n'
    '- NO external links\n'
    '- Minimal promotional language\n'
    '- Tone: human, reflective, conversational\n'
    '- Platform: LinkedIn (B2B audience)'
)

add_body(
    'These constraints are directly informed by the Engagement Prediction module\'s findings: '
    'optimal post length is 100\u2013200 words, external links incur algorithm penalties, and authentic '
    'tone drives higher comment engagement.'
)

add_heading_styled('4.3 User Prompt Builder', level=2)
add_body(
    'The user prompt builder accepts four parameters to create contextualized generation requests:'
)
params = [
    ('topic: ', 'The trending topic from Module 1 (e.g., "AI regulation in fintech")'),
    ('personal_context: ', 'User\'s professional background for authentic voice'),
    ('insight: ', 'Core learning or perspective shift to anchor the post'),
    ('hook_style: ', 'Opening sentence pattern (contrarian, transformation, hidden insight)'),
]
for bold_part, rest in params:
    add_bullet(rest, bold_prefix=bold_part)

add_body(
    'The prompt enforces LinkedIn-specific constraints: 100\u2013200 words, no external links, no promotional '
    'language, short paragraphs, and ending with a thoughtful question to drive comments.'
)

add_heading_styled('4.4 Multi-Variant Generation', level=2)
add_body(
    'For each trending topic, the system generates three post variants using different hook styles:'
)
hooks = [
    ('Contrarian: ', '"Everyone thinks X, but actually Y..." \u2013 Challenges conventional wisdom'),
    ('Personal Transformation: ', '"I used to believe X, now I realize Y..." \u2013 Authentic growth narrative'),
    ('Hidden Insight: ', '"Nobody talks about X..." \u2013 Reveals non-obvious perspective'),
]
for bold_part, rest in hooks:
    add_bullet(rest, bold_prefix=bold_part)

add_body(
    'This multi-variant approach enables A/B testing and allows the Engagement Prediction module '
    '(Module 3) to score each variant, recommending the highest-predicted performer.'
)

add_heading_styled('4.5 Sample Output', level=2)
add_body('For the topic "AI regulation in fintech" with a personal transformation hook:')
add_code_block(
    'I used to believe that regulation in fintech was a roadblock\n'
    'to innovation. I thought it stifled creativity and slowed\n'
    'down progress.\n\n'
    'But after spending years as a data professional in financial\n'
    'services, my perspective has shifted dramatically. I\'ve\n'
    'witnessed firsthand how regulations can actually reshape\n'
    'innovation, providing a framework that encourages new ideas\n'
    'while ensuring consumer protection.\n\n'
    'For instance, when GDPR came into play, many saw it as a\n'
    'hurdle. However, it pushed us to rethink data privacy and\n'
    'security, leading to more robust and trustworthy financial\n'
    'products.\n\n'
    'How do you see regulation influencing innovation in your field?'
)

add_heading_styled('4.6 Current Status and Future Work', level=2)
add_body(
    'The Post Generation module is currently implemented as a proof-of-concept notebook. Remaining work '
    'includes integration with Module 1\'s trend output as direct input, connection to Module 3 for '
    'engagement scoring of variants, packaging as a Streamlit agent, and adding user profile-based '
    'tone customization.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. MODULE 3: ENGAGEMENT PREDICTION
# ═══════════════════════════════════════════════════════════════
add_heading_styled('5. Module 3: Engagement Prediction', level=1)

add_heading_styled('5.1 Abstract and Objectives', level=2)
add_body(
    'This module presents a comprehensive machine learning approach to predicting engagement metrics '
    '(reactions and comments) for LinkedIn posts. Working with a dataset of 31,996 posts from 69 verified '
    'LinkedIn influencers, we developed and validated predictive models achieving R\u00b2 scores of 0.5903 '
    'for reactions and 0.5280 for comments.'
)
add_body(
    'The study encompasses data cleaning, extensive feature engineering (85 features across 9 categories), '
    'multi-algorithm evaluation, and rigorous validation to ensure production readiness.'
)

add_heading_styled('5.2 Data Collection and Preprocessing', level=2)
add_body(
    'The dataset comprises LinkedIn posts from 69 verified influencers collected through web scraping '
    'of public profiles (Kaggle). Initial collection yielded 34,012 posts with 19 features.'
)
add_body(
    'To improve model robustness and reduce bias from the original influencer-heavy dataset, we augmented '
    'the training data with additional real posts scraped directly from LinkedIn. This expanded dataset '
    'provides broader coverage of posting styles, follower ranges, and engagement levels. Crucially, we '
    'addressed the strong confounding effect of follower count on raw engagement numbers. Follower counts '
    'were log-transformed to compress the heavily right-skewed distribution (ranging from a few hundred to '
    'over 100,000), and engagement scores were normalised by dividing by the number of followers. This '
    'per-follower normalisation ensures the model learns content quality signals rather than simply '
    'predicting that larger audiences produce higher absolute engagement. The combination of log-transformed '
    'follower counts and follower-normalised engagement targets significantly reduced the dominance of '
    'influencer profile features in the model, allowing content-quality features (readability, sentiment, '
    'hook patterns) to contribute more meaningfully to predictions.'
)

add_heading_styled('Data Quality Issues and Treatment', level=3)
add_table(
    ['Issue', 'Scope', 'Treatment'],
    [
        ['Missing content', '2,016 posts (5.93%)', 'Removed rows'],
        ['Missing views', '100%', 'Column excluded'],
        ['Extreme outliers', 'Max reactions: 391,498 vs median 38', 'Winsorization at 99th percentile'],
        ['Non-numeric followers', '42 profiles (0.13%)', 'Median imputation'],
        ['Duplicate content', '757 posts (2.2%)', 'Retained (legitimate cross-author content)'],
    ]
)

add_heading_styled('Text Preprocessing Pipeline', level=3)
preprocessing = [
    ('URL Extraction: ', 'External links extracted (20% of posts) and replaced with [URL] placeholder'),
    ('Mention Extraction: ', 'User mentions replaced with [MENTION] tokens (2.9% of posts)'),
    ('Hashtag Treatment: ', 'Extracted for counting but preserved in text (50.9% of posts, avg 4.83)'),
    ('Emoji Processing: ', 'Extracted, counted (6.9% of posts), then removed from clean text'),
    ('Normalization: ', 'Lowercase conversion, whitespace standardization, special character retention'),
]
for bold_part, rest in preprocessing:
    add_bullet(rest, bold_prefix=bold_part)

add_heading_styled('5.3 Feature Engineering (85 Features)', level=2)
add_body(
    'Feature engineering represents the most critical phase of model development, transforming raw '
    'content into machine-learning-ready representations. We developed 85 features spanning 9 categories.'
)

add_table(
    ['Category', 'Count', 'Key Features'],
    [
        ['Base Formula', '15', 'Content length score, hook patterns (9 types), power patterns (15), media score, link penalty'],
        ['NLP \u2013 Sentiment', '5', 'VADER positive/negative/neutral/compound scores, sentiment category'],
        ['NLP \u2013 Named Entities', '12', 'Entity counts by type (PERSON, ORG, GPE, DATE, etc.), presence flags'],
        ['NLP \u2013 Readability', '5', 'Flesch Reading Ease, Flesch-Kincaid, SMOG, Gunning Fog, ARI'],
        ['NLP \u2013 Text Stats', '8', 'Sentence count, avg sentence length, lexical diversity, syllable metrics'],
        ['NLP \u2013 Stylistic', '13', 'Question/exclamation marks, ALL CAPS, quotes, bullet points, emojis'],
        ['Topic Classification', '7', 'Tech/AI, Business, Career, Leadership, Personal Dev, Finance + composites'],
        ['Influencer Profile', '12', 'Avg/median engagement, consistency score, post count, total engagement'],
        ['Derived/Interaction', '13', 'Engagement ratios, density metrics, sentiment\u00d7readability, media\u00d7sentiment'],
    ]
)

add_heading_styled('Data Leakage Discovery and Resolution', level=3)
add_body(
    'During initial model development (Version 1), we discovered that 6 derived features contained '
    'target information (e.g., reactions_per_word, comment_to_reaction_ratio). These features inflated '
    'performance artificially (R\u00b2 > 0.99) but would fail in production. All six were removed in '
    'Version 2, reducing R\u00b2 to realistic levels (0.59 reactions, 0.53 comments).'
)

add_heading_styled('5.4 Model Training and Selection', level=2)
add_body(
    'Five algorithm families were evaluated using an 80-20 train-test split (25,596 training / '
    '6,400 test posts) with StandardScaler normalization:'
)

add_heading_styled('Reactions Prediction Performance', level=3)
add_table(
    ['Model', 'R\u00b2 Score', 'MAE', 'RMSE', 'sMAPE (%)'],
    [
        ['Random Forest (Winner)', '0.5903', '191.68', '601.68', '74.16'],
        ['LightGBM', '0.5816', '197.25', '608.12', '76.43'],
        ['XGBoost', '0.5718', '204.33', '615.27', '79.18'],
        ['Ridge', '0.5096', '242.56', '658.34', '88.92'],
        ['Linear Regression', '0.5095', '242.57', '658.35', '88.93'],
    ]
)

add_heading_styled('Comments Prediction Performance', level=3)
add_table(
    ['Model', 'R\u00b2 Score', 'MAE', 'RMSE', 'sMAPE (%)'],
    [
        ['LightGBM (Winner)', '0.5280', '15.26', '36.36', '117.08'],
        ['Random Forest', '0.5250', '15.00', '36.48', '109.90'],
        ['XGBoost', '0.5200', '15.22', '36.67', '114.46'],
        ['Ridge', '0.4077', '19.44', '40.73', '129.71'],
        ['Linear Regression', '0.4076', '19.44', '40.74', '129.72'],
    ]
)

add_heading_styled('Cross-Validation Results', level=3)
add_body(
    'Reactions (Random Forest): CV Mean = 0.6118 \u00b1 0.0600, Test R\u00b2 = 0.5903 \u2013 No overfitting confirmed. '
    'Comments (LightGBM): CV Mean = 0.5496 \u00b1 0.0643, Test R\u00b2 = 0.5280 \u2013 Stable generalization confirmed.'
)

add_heading_styled('5.5 Results and Feature Importance', level=2)

add_heading_styled('Top Predictors of Reactions', level=3)
add_table(
    ['Rank', 'Feature', 'Importance', 'Category'],
    [
        ['1', 'influencer_avg_engagement', '36.2%', 'Influencer'],
        ['2', 'influencer_total_engagement', '29.6%', 'Influencer'],
        ['3', 'text_difficult_words_ratio', '3.5%', 'NLP'],
        ['4', 'influencer_post_count', '2.9%', 'Influencer'],
        ['5', 'influencer_consistency_reactions', '2.4%', 'Influencer'],
    ]
)

add_body(
    'Critical Finding: The top 5 features (all influencer-related) account for 74.6% of total predictive '
    'power, revealing that reactions are primarily determined by WHO posts content rather than WHAT '
    'content is posted.'
)

add_heading_styled('Top Predictors of Comments', level=3)
add_table(
    ['Rank', 'Feature', 'Importance', 'Category'],
    [
        ['1', 'influencer_avg_engagement', '32.1%', 'Influencer'],
        ['2', 'text_difficult_words_ratio', '14.4%', 'NLP'],
        ['3', 'influencer_total_engagement', '13.6%', 'Influencer'],
        ['4', 'readability_ari', '13.5%', 'NLP'],
        ['5', 'text_avg_sentence_length', '13.2%', 'NLP'],
    ]
)

add_body(
    'Critical Finding: Unlike reactions (68% influencer dominance), comments show more balanced '
    'attribution: Influencer features 45.7%, NLP/Content features 54.3%. This confirms that content '
    'quality drives active engagement (comments) more than passive engagement (reactions).'
)

add_heading_styled('Feature Category Comparison', level=3)
add_table(
    ['Category', 'Reactions Importance', 'Comments Importance', 'Difference'],
    [
        ['Influencer Profile', '74.6%', '45.7%', '+28.9% (Reactions)'],
        ['NLP/Text Quality', '8.3%', '38.1%', '+29.8% (Comments)'],
        ['Sentiment', '2.1%', '12.3%', '+10.2% (Comments)'],
        ['Base Formula', '5.8%', '11.9%', '+6.1% (Comments)'],
        ['Media/Visuals', '4.2%', '2.9%', '+1.3% (Reactions)'],
        ['Derived Features', '3.5%', '12.5%', '+9.0% (Comments)'],
    ]
)

add_heading_styled('5.6 Validation and Robustness', level=2)
validation_items = [
    ('Residual Analysis: ', 'Residuals approximately normally distributed and centered at zero. '
     'Model underpredicts high-engagement posts (>3,000 reactions) due to unpredictable viral factors.'),
    ('Edge Case Testing: ', 'All 34 edge cases passed, including zero engagement posts, missing features, '
     'extreme values, and mega-influencer predictions.'),
    ('Inference Performance: ', 'Reactions model <15ms, Comments model <20ms, Feature engineering ~50ms. '
     'Total latency ~70ms end-to-end per post. Memory footprint <50 MB.'),
]
for bold_part, rest in validation_items:
    add_bullet(rest, bold_prefix=bold_part)

add_heading_styled('5.7 Key Findings', level=2)
findings = [
    ('Finding 1 \u2013 Influencer Dominance: ', 'Historical author metrics account for 36\u201375% of engagement variance. '
     'Reactions are primarily author-driven; comments are content-driven.'),
    ('Finding 2 \u2013 Dual-Process Engagement: ', 'Reactions (passive, one-click) are driven by source trust. '
     'Comments (active, typing effort) require compelling content regardless of author fame.'),
    ('Finding 3 \u2013 Sufficient Accuracy: ', 'R\u00b2 = 0.59 (reactions) and 0.53 (comments) exceed targets, enabling '
     'reliable A/B testing of post variants and engagement band estimation.'),
    ('Finding 4 \u2013 Data Leakage Lesson: ', 'Initial leaky models achieved R\u00b2 > 0.99. The 90% performance drop '
     'after cleanup demonstrates the importance of systematic leakage audits.'),
]
for bold_part, rest in findings:
    add_bullet(rest, bold_prefix=bold_part)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. MODULE 4: VISUAL GENERATION & PUBLISHING SUPPORT
# ═══════════════════════════════════════════════════════════════
add_heading_styled('6. Module 4: Visual Generation & Publishing Support', level=1)

add_heading_styled('6.1 Architecture Overview', level=2)
add_body(
    'The Visual Generation module provides automated visual content creation to accompany LinkedIn posts. '
    'It implements a unified CLI service (visual_service.py, 148 lines) that integrates three sub-components: '
    'an LLM classification layer, a diagram renderer, and an image generator.'
)
add_body(
    'The service follows a classify-then-route pattern: user input is first analyzed by an LLM to determine '
    'whether the visual need is best served by a diagram (flowcharts, architecture diagrams, sequences) or '
    'a photographic/artistic image. Based on this classification, the request is routed to the appropriate '
    'generator.'
)

add_heading_styled('6.2 LLM Classification Layer', level=2)
add_body(
    'The classify_and_generate() function (from llm/classify_llm.py) processes user input through GPT-4o '
    'and returns a structured JSON response indicating the visual type:'
)
add_bullet('type: "diagram" \u2013 Routes to Mermaid renderer with generated Mermaid code', bold_prefix='Diagram Classification: ')
add_bullet('type: "image" \u2013 Routes to Stable Diffusion with an optimized image prompt', bold_prefix='Image Classification: ')
add_bullet('type: "none" \u2013 No visual content needed for this input', bold_prefix='No Visual: ')

add_body(
    'The LLM not only classifies but also generates the appropriate content: for diagrams, it produces '
    'valid Mermaid.js syntax; for images, it crafts an optimized prompt for the image generation model.'
)

add_heading_styled('6.3 Diagram Generation (Mermaid)', level=2)
add_body(
    'The render_mermaid() function (from diagram/render_mermaid.py) takes Mermaid.js code and renders '
    'it to a PNG image file. Mermaid supports multiple diagram types relevant to professional LinkedIn content:'
)
diagram_types = [
    'Flowcharts \u2013 Process workflows, decision trees',
    'Sequence diagrams \u2013 API interactions, system communication',
    'Architecture diagrams \u2013 System design, infrastructure layouts',
    'Gantt charts \u2013 Project timelines, roadmaps',
    'Mind maps \u2013 Concept relationships, brainstorming',
]
for d in diagram_types:
    add_bullet(d)

add_body(
    'Error handling includes logging of failed Mermaid source code to timestamped .mmd files for '
    'debugging, ensuring that LLM-generated diagram code can be inspected and corrected.'
)

add_heading_styled('6.4 Image Generation (Stable Diffusion)', level=2)
add_body(
    'The generate_image() function (from image/generate_image.py) uses Stable Diffusion to create '
    'photographic or artistic images from text prompts. The LLM classification layer optimizes the prompt '
    'for image generation quality before passing it to the model.'
)
add_body(
    'Generated images are saved as PNG files in the service directory, with paths returned as JSON '
    'for downstream consumption by the Streamlit application.'
)

add_heading_styled('6.5 Service Orchestration', level=2)
add_body(
    'The main() function in visual_service.py orchestrates the full pipeline:'
)

orchestration_steps = [
    'Accept user input via CLI argument',
    'Call LLM classification to determine visual type and generate content',
    'Route to appropriate generator (Mermaid or Stable Diffusion)',
    'Return JSON output with type, format, and file path',
    'Handle errors gracefully with logging and fallback Mermaid source persistence',
]
for i, step in enumerate(orchestration_steps, 1):
    add_bullet(step, bold_prefix=f'{i}. ')

add_body(
    'The service uses structured logging (visual_service.log) with both console (INFO+) and file '
    '(DEBUG) handlers for traceability. All LLM responses and generated prompts are logged for '
    'debugging and quality assessment.'
)

add_heading_styled('6.6 Key Design Decisions', level=2)
decisions = [
    ('Direct function imports: ', 'Sub-modules are imported at the module level rather than spawning '
     'subprocesses, reducing latency and enabling better error handling.'),
    ('Lazy file handler: ', 'Log file handler is attached only when main() runs, preventing disk writes '
     'during module import (important for testing).'),
    ('Failed diagram persistence: ', 'When Mermaid rendering fails, the source code is saved to a '
     'timestamped .mmd file for manual inspection and LLM prompt debugging.'),
    ('JSON output contract: ', 'All outputs follow a consistent JSON schema (type, format, path) '
     'for seamless integration with the Streamlit frontend.'),
]
for bold_part, rest in decisions:
    add_bullet(rest, bold_prefix=bold_part)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. INTEGRATION ARCHITECTURE
# ═══════════════════════════════════════════════════════════════
add_heading_styled('7. Integration Architecture', level=1)

add_body(
    'The four modules are designed to operate as a sequential pipeline with well-defined data contracts '
    'between each stage. The integration architecture ensures that each module can be developed, tested, '
    'and enhanced independently.'
)

add_heading_styled('7.1 Data Flow Between Modules', level=2)
add_table(
    ['From', 'To', 'Data Contract', 'Format'],
    [
        ['Module 1 (Trends)', 'Module 2 (Posts)', 'Trending topic + search context + rising queries', 'JSON dict'],
        ['Module 2 (Posts)', 'Module 3 (Engagement)', 'Generated post text + metadata (length, hooks, media)', 'Feature vector (85 dims)'],
        ['Module 3 (Engagement)', 'User/Streamlit', 'Predicted reactions + comments per variant', 'JSON scores'],
        ['Module 2 (Posts)', 'Module 4 (Visuals)', 'Post content or user visual request', 'Text string'],
        ['Module 4 (Visuals)', 'User/Streamlit', 'Generated image/diagram path', 'JSON with file path'],
    ]
)

add_heading_styled('7.2 Planned Streamlit Application', level=2)
add_body(
    'The Streamlit application will serve as the unified frontend, providing:'
)
streamlit_features = [
    'LinkedIn bio input and domain selection interface',
    'Real-time trend discovery with interactive topic selection',
    'Post generation with multiple variants displayed side-by-side',
    'Engagement prediction scores for each variant with visual comparison',
    'Visual generation triggered by post content or user request',
    'Final review and one-click publishing workflow',
]
for f in streamlit_features:
    add_bullet(f)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 8. CURRENT STATUS SUMMARY
# ═══════════════════════════════════════════════════════════════
add_heading_styled('8. Current Status Summary', level=1)

add_table(
    ['Module', 'Component', 'Status', 'Completion'],
    [
        ['1. Trend Identification', 'Core algorithm (4 iterations)', 'Complete', '100%'],
        ['1. Trend Identification', 'Google Trends integration', 'Complete', '100%'],
        ['1. Trend Identification', 'GPT-4o prompt engineering', 'Complete', '100%'],
        ['1. Trend Identification', 'Interactive topic selection', 'Complete', '100%'],
        ['1. Trend Identification', 'Hallucination guards', 'Complete', '100%'],
        ['2. Post Generation', 'Master system prompt', 'Complete', '100%'],
        ['2. Post Generation', 'Multi-variant generation', 'Complete', '100%'],
        ['2. Post Generation', 'Streamlit agent integration', 'Pending', '20%'],
        ['3. Engagement Prediction', 'Data collection & cleaning', 'Complete', '100%'],
        ['3. Engagement Prediction', 'Feature engineering (85 features)', 'Complete', '100%'],
        ['3. Engagement Prediction', 'Model training & validation', 'Complete', '100%'],
        ['3. Engagement Prediction', 'Production deployment prep', 'In Progress', '80%'],
        ['4. Visual Generation', 'LLM classification layer', 'Complete', '100%'],
        ['4. Visual Generation', 'Mermaid diagram rendering', 'Complete', '100%'],
        ['4. Visual Generation', 'Stable Diffusion integration', 'Complete', '100%'],
        ['4. Visual Generation', 'CLI service orchestration', 'Complete', '100%'],
        ['Integration', 'Unit testing (all modules)', 'Pending', '0%'],
        ['Integration', 'Integration testing', 'Pending', '0%'],
        ['Integration', 'Streamlit application', 'Pending', '10%'],
        ['Integration', 'End-to-end pipeline', 'Pending', '5%'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 9. PENDING WORK AND NEXT STEPS
# ═══════════════════════════════════════════════════════════════
add_heading_styled('9. Pending Work and Next Steps', level=1)

add_heading_styled('9.1 Unit Testing', level=2)
add_body('Unit tests are needed across all four modules:')
unit_tests = [
    ('Module 1: ', 'Test GPT-4o topic extraction with mock bios, Google Trends API mocking, cache '
     'validation logic, interactive selection loop, version hallucination guard effectiveness.'),
    ('Module 2: ', 'Test prompt builder parameter injection, system prompt constraint enforcement, '
     'word count validation, hook style variation, edge cases (empty topic, very long bios).'),
    ('Module 3: ', 'Test feature engineering pipeline with known inputs, model prediction consistency, '
     'edge case handling (zero engagement, missing features, extreme values), serialization/deserialization '
     'of model artifacts.'),
    ('Module 4: ', 'Test LLM classification routing, Mermaid code validation, image generation with mock '
     'models, error handling paths, JSON output contract compliance.'),
]
for bold_part, rest in unit_tests:
    add_bullet(rest, bold_prefix=bold_part)

add_heading_styled('9.2 Integration Testing', level=2)
add_body('End-to-end integration tests verifying data flow between modules:')
integration_tests = [
    'Module 1 \u2192 Module 2: Trend output correctly consumed by post generator',
    'Module 2 \u2192 Module 3: Generated posts correctly featurized and scored',
    'Module 2 \u2192 Module 4: Post content triggers appropriate visual generation',
    'Full pipeline: Bio input \u2192 Trends \u2192 Posts \u2192 Scores \u2192 Visuals \u2192 Final output',
    'Error propagation: Upstream failures handled gracefully by downstream modules',
]
for t in integration_tests:
    add_bullet(t)

add_heading_styled('9.3 Streamlit Application Development', level=2)
add_body('The Streamlit application requires:')
streamlit_tasks = [
    'Page layout design with multi-step workflow wizard',
    'LinkedIn bio input form with validation',
    'Real-time trend display with interactive selection widgets',
    'Post variant comparison view with engagement score overlay',
    'Visual generation trigger and display panel',
    'Session state management across pipeline stages',
    'Error handling and user feedback mechanisms',
]
for t in streamlit_tasks:
    add_bullet(t)

add_heading_styled('9.4 Future Enhancements', level=2)
future = [
    'Geographic customization for trend identification (currently hardcoded to US)',
    'Multi-profile batch analysis for enterprise users',
    'Transformer-based models (BERT/GPT fine-tuning) for engagement prediction',
    'SHAP values for instance-level engagement explanations',
    'Content-only models for cold-start users without historical data',
    'Scheduled execution and database persistence for trend monitoring',
    'LinkedIn API integration for direct publishing',
]
for f in future:
    add_bullet(f)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 10. CONCLUSION
# ═══════════════════════════════════════════════════════════════
add_heading_styled('10. Conclusion', level=1)

add_body(
    'TrendPilot has achieved substantial progress across all four modules of the AI-powered LinkedIn '
    'content strategy pipeline. The project demonstrates a coherent, end-to-end approach to automating '
    'professional content creation \u2013 from trend discovery through engagement-optimized post generation '
    'to visual content creation.'
)

add_body(
    'The Trending Topic Identification module (Module 1) evolved through four iterations, each addressing '
    'specific shortcomings: from generic news-based results to domain-aware LLM extraction, from ambiguous '
    'entity recognition to precise technical topic identification, and from one-shot execution to interactive '
    'user-driven exploration. The final system combines GPT-4o\'s contextual understanding with Google Trends\' '
    'real-time search data to produce actionable, timely post recommendations.'
)

add_body(
    'The Post Generation module (Module 2) implements a sophisticated two-prompt architecture informed by '
    'engagement research findings. By generating multiple variants with different hook styles, the system '
    'enables data-driven content selection through the Engagement Prediction module.'
)

add_body(
    'The Engagement Prediction module (Module 3) represents the most rigorous component, with 85 engineered '
    'features, five algorithm evaluations, systematic data leakage prevention, and comprehensive validation. '
    'The discovery that reactions are author-driven (75%) while comments are content-driven (54%) provides '
    'both theoretical insight and practical guidance for content optimization.'
)

add_body(
    'The Visual Generation module (Module 4) completes the pipeline with intelligent visual content creation, '
    'using LLM classification to route between diagram and image generation based on content context.'
)

add_body(
    'The remaining work \u2013 unit testing, integration testing, and Streamlit application development \u2013 '
    'represents the final integration phase that will transform these individual modules into a cohesive, '
    'user-facing application. The modular architecture ensures this integration can proceed incrementally '
    'without disrupting the proven core algorithms.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 11. REFERENCES
# ═══════════════════════════════════════════════════════════════
add_heading_styled('11. References', level=1)

references = [
    'Chua, T. H. H., & Banerjee, S. (2015). Understanding user engagement in social media: A study of Facebook likes. Journal of Interactive Marketing, 31, 13-25.',
    'De Vries, L., Gensler, S., & Leeflang, P. S. H. (2012). Popularity of brand posts on brand fan pages: An investigation of the effects of social media marketing. Journal of Interactive Marketing, 26(2), 83-91.',
    'Hoffman, D. L., & Fodor, M. (2010). Can you measure the ROI of your social media marketing? MIT Sloan Management Review, 52(1), 41-49.',
    'Kaufman, S., Rosset, S., Perlich, C., & Stitelman, O. (2012). Leakage in data mining: Formulation, detection, and avoidance. ACM Transactions on Knowledge Discovery from Data, 6(4), 1-21.',
    'Suh, B., Hong, L., Pirolli, P., & Chi, E. H. (2010). Want to be retweeted? Large scale analytics on factors impacting retweet in Twitter network. IEEE International Conference on Social Computing.',
    'LinkedIn Influencers Dataset: https://www.kaggle.com/datasets/shreyasajal/linkedin-influencers-data',
]
for i, ref in enumerate(references, 1):
    p = doc.add_paragraph(f'[{i}] {ref}')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)

# ── Save ──
output_path = os.path.join(os.path.dirname(__file__), 'TrendPilot_Comprehensive_Report.docx')
doc.save(output_path)
print(f'Report saved to: {output_path}')
