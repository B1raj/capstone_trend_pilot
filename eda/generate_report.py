"""
Generate comprehensive Word document report from LinkedIn EDA notebook findings
With embedded diagrams and detailed explanations
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Base path for images
BASE_PATH = '/Users/BirajMishra/work/playground/ms/capstone/capstone_trend_pilot/eda'

# Create document
doc = Document()

# Helper function to add horizontal line
def add_horizontal_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_break()

# Helper function to add image with caption
def add_image_with_caption(doc, image_path, caption, width=6):
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(width))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        caption_para = doc.add_paragraph()
        caption_run = caption_para.add_run(f'Figure: {caption}')
        caption_run.italic = True
        caption_run.font.size = Pt(10)
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Spacing
    else:
        doc.add_paragraph(f'[Image not found: {image_path}]')

# ============================================================================
# TITLE PAGE
# ============================================================================
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_heading('LinkedIn Influencer Data', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('Exploratory Data Analysis Report', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()

# Project info
info_para = doc.add_paragraph()
info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_para.add_run('Capstone Project: TrendPilot\n\n').bold = True
info_para.add_run('A System for Creating Engaging LinkedIn Posts\n\n')
info_para.add_run('Data Analysis Phase\n\n')

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# ============================================================================
# TABLE OF CONTENTS
# ============================================================================
doc.add_page_break()
doc.add_heading('Table of Contents', level=1)

toc_items = [
    ('1. Executive Summary', 3),
    ('2. Introduction and Objectives', 4),
    ('3. Dataset Overview', 5),
    ('4. Data Gap Analysis - Missing Values', 6),
    ('5. Data Quality Assessment', 9),
    ('6. Distribution Analysis of Key Metrics', 12),
    ('7. Content Analysis', 15),
    ('8. Engagement Analysis', 18),
    ('9. Correlation Analysis', 22),
    ('10. Top Performers Analysis', 24),
    ('11. Data Gap Risk Assessment', 26),
    ('12. Impact on Final Project', 28),
    ('13. Recommendations', 31),
    ('14. Conclusion', 33),
]

for item, page in toc_items:
    toc_para = doc.add_paragraph()
    toc_para.add_run(item)
    toc_para.add_run('\t' * 8 + str(page))

# ============================================================================
# 1. EXECUTIVE SUMMARY
# ============================================================================
doc.add_page_break()
doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    'This comprehensive Exploratory Data Analysis (EDA) report examines a dataset of LinkedIn posts '
    'from 69 influential professionals, comprising 34,012 individual posts. The analysis was conducted '
    'to support the development of TrendPilot, a system designed to help users create engaging LinkedIn '
    'posts optimized for maximum likes (reactions) and comments.'
)

doc.add_paragraph(
    'The primary objectives of this EDA were to: (1) identify data gaps and missing values that could '
    'impact model development, (2) assess the quality and reliability of the available data, (3) understand '
    'the distribution and relationships between key engagement metrics, and (4) evaluate how data limitations '
    'might affect the final project deliverables.'
)

doc.add_heading('Key Findings at a Glance', level=2)

key_findings = [
    ('Data Completeness', 'The dataset achieves 86.7% overall completeness and 95.4% completeness for core '
     'features essential to engagement prediction. This indicates the data is suitable for building '
     'machine learning models.'),
    ('Critical Data Gap', 'Views data is 100% missing across all records. This is a significant limitation '
     'as it prevents optimization for post reach and virality metrics.'),
    ('Engagement Metrics', 'Reactions and comments data is 100% complete, providing a solid foundation for '
     'building engagement prediction models. These metrics show a strong positive correlation (r=0.823).'),
    ('Content Availability', '94.1% of posts have content text available, enabling robust natural language '
     'processing (NLP) analysis for content optimization features.'),
    ('Media Impact', 'Video and image content significantly outperform articles in engagement metrics, with '
     'videos generating an average of 866 reactions compared to 161 for articles.'),
    ('Temporal Limitation', 'Timestamps are stored in relative format ("1 day ago", "2 weeks ago") rather than '
     'absolute dates, preventing optimal posting time recommendations.'),
]

for title_text, description in key_findings:
    para = doc.add_paragraph()
    para.add_run(f'{title_text}: ').bold = True
    para.add_run(description)

doc.add_heading('Data Quality Assessment Summary', level=2)

summary_table = doc.add_table(rows=4, cols=2)
summary_table.style = 'Table Grid'
summary_data = [
    ('Overall Data Completeness', '86.7%'),
    ('Core Features Completeness', '95.4%'),
    ('Quality Assessment', 'GOOD'),
    ('Suitability for ML', 'SUITABLE'),
]
for i, (metric, value) in enumerate(summary_data):
    summary_table.rows[i].cells[0].text = metric
    summary_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    summary_table.rows[i].cells[1].text = value

# ============================================================================
# 2. INTRODUCTION AND OBJECTIVES
# ============================================================================
doc.add_page_break()
doc.add_heading('2. Introduction and Objectives', level=1)

doc.add_heading('2.1 Background', level=2)
doc.add_paragraph(
    'LinkedIn has emerged as the premier professional networking platform, with over 900 million users '
    'worldwide. For professionals, thought leaders, and businesses, creating engaging content on LinkedIn '
    'is crucial for building personal brands, establishing thought leadership, and driving business outcomes. '
    'However, the factors that contribute to high-engagement posts are not always intuitive, and users often '
    'struggle to optimize their content for maximum impact.'
)

doc.add_paragraph(
    'This EDA supports the development of TrendPilot, an intelligent system designed to analyze patterns '
    'in successful LinkedIn posts and provide data-driven recommendations for content creation. By studying '
    'posts from verified influencers who consistently achieve high engagement, we can extract insights that '
    'will benefit everyday LinkedIn users.'
)

doc.add_heading('2.2 Analysis Objectives', level=2)
doc.add_paragraph('This exploratory data analysis was conducted with the following specific objectives:')

objectives = [
    ('Identify Data Gaps', 'Systematically catalog missing values, incomplete records, and data quality '
     'issues that could affect downstream analysis and model development.'),
    ('Assess Data Quality', 'Evaluate the reliability, consistency, and accuracy of the collected data '
     'through statistical analysis and anomaly detection.'),
    ('Understand Engagement Patterns', 'Analyze the distribution of engagement metrics (reactions, comments, '
     'views) and identify patterns associated with high-performing content.'),
    ('Evaluate Feature Relationships', 'Examine correlations between content characteristics (length, hashtags, '
     'media type) and engagement outcomes.'),
    ('Impact Assessment', 'Determine how identified data gaps and quality issues will affect the feasibility '
     'and scope of planned features in the final system.'),
]

for obj_title, obj_desc in objectives:
    para = doc.add_paragraph(style='List Number')
    para.add_run(f'{obj_title}: ').bold = True
    para.add_run(obj_desc)

doc.add_heading('2.3 Methodology', level=2)
doc.add_paragraph(
    'The analysis was conducted using Python with the following key libraries: Pandas for data manipulation, '
    'NumPy for numerical operations, Matplotlib and Seaborn for static visualizations, and Plotly for '
    'interactive dashboards. The analysis followed a structured approach:'
)

methodology_steps = [
    'Data loading and initial inspection',
    'Missing value analysis and visualization',
    'Data type verification and conversion',
    'Statistical summary generation',
    'Distribution analysis with appropriate transformations',
    'Correlation analysis between features',
    'Segmented analysis by media type and influencer',
    'Risk assessment and impact evaluation',
]
for step in methodology_steps:
    doc.add_paragraph(step, style='List Bullet')

# ============================================================================
# 3. DATASET OVERVIEW
# ============================================================================
doc.add_page_break()
doc.add_heading('3. Dataset Overview', level=1)

doc.add_heading('3.1 Data Source and Collection', level=2)
doc.add_paragraph(
    'The dataset comprises LinkedIn posts collected from 69 verified influencers across various industries '
    'and professional domains. These influencers were selected based on their established presence on the '
    'platform, consistent posting activity, and demonstrated ability to generate engagement. The data '
    'includes both profile-level information (name, headline, location, followers) and post-level details '
    '(content, media, hashtags, engagement metrics).'
)

doc.add_heading('3.2 Dataset Dimensions', level=2)

dims_table = doc.add_table(rows=4, cols=2)
dims_table.style = 'Table Grid'
dims_data = [
    ('Total Records (Posts)', '34,012'),
    ('Total Features (Columns)', '19'),
    ('Unique Influencers', '69'),
    ('Average Posts per Influencer', '493'),
]
for i, (metric, value) in enumerate(dims_data):
    dims_table.rows[i].cells[0].text = metric
    dims_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    dims_table.rows[i].cells[1].text = value

doc.add_paragraph()

doc.add_heading('3.3 Feature Descriptions', level=2)
doc.add_paragraph(
    'The dataset contains 19 features capturing various aspects of influencer profiles and their posts. '
    'Understanding each feature is crucial for effective analysis and feature engineering.'
)

doc.add_heading('Profile-Level Features', level=3)
profile_features = [
    ('slno', 'int64', 'Sequential identifier for each record'),
    ('name', 'object', 'Full name of the LinkedIn influencer'),
    ('headline', 'object', 'Professional headline displayed on the profile'),
    ('location', 'object', 'Geographic location of the influencer'),
    ('followers', 'float64', 'Total number of followers for the profile'),
    ('connections', 'object', 'Number of connections (often shows "500+" for max)'),
    ('about', 'object', 'Biography/about section from the profile'),
]

profile_table = doc.add_table(rows=len(profile_features)+1, cols=3)
profile_table.style = 'Table Grid'
profile_headers = ['Feature', 'Data Type', 'Description']
for i, h in enumerate(profile_headers):
    profile_table.rows[0].cells[i].text = h
    profile_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
for i, (feat, dtype, desc) in enumerate(profile_features, 1):
    profile_table.rows[i].cells[0].text = feat
    profile_table.rows[i].cells[1].text = dtype
    profile_table.rows[i].cells[2].text = desc

doc.add_paragraph()

doc.add_heading('Post-Level Features', level=3)
post_features = [
    ('time_spent', 'object', 'Relative time since post was published (e.g., "1 day ago")'),
    ('content', 'object', 'Text content of the LinkedIn post'),
    ('content_links', 'object', 'Hyperlinks mentioned in the post content'),
    ('media_type', 'object', 'Type of media attached (article, image, video, etc.)'),
    ('media_url', 'object', 'URL(s) of attached media'),
    ('num_hashtags', 'int64', 'Count of hashtags used in the post'),
    ('hashtag_followers', 'int64', 'Aggregate followers of hashtags used'),
    ('hashtags', 'object', 'List of hashtags with their URLs'),
]

post_table = doc.add_table(rows=len(post_features)+1, cols=3)
post_table.style = 'Table Grid'
for i, h in enumerate(profile_headers):
    post_table.rows[0].cells[i].text = h
    post_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
for i, (feat, dtype, desc) in enumerate(post_features, 1):
    post_table.rows[i].cells[0].text = feat
    post_table.rows[i].cells[1].text = dtype
    post_table.rows[i].cells[2].text = desc

doc.add_paragraph()

doc.add_heading('Engagement Metrics', level=3)
engagement_features = [
    ('reactions', 'int64', 'Total reactions (likes, celebrates, etc.) on the post'),
    ('comments', 'int64', 'Total number of comments on the post'),
    ('views', 'float64', 'Number of views/impressions (NOTE: 100% missing)'),
    ('votes', 'object', 'Poll votes if the post contains a poll'),
]

eng_table = doc.add_table(rows=len(engagement_features)+1, cols=3)
eng_table.style = 'Table Grid'
for i, h in enumerate(profile_headers):
    eng_table.rows[0].cells[i].text = h
    eng_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
for i, (feat, dtype, desc) in enumerate(engagement_features, 1):
    eng_table.rows[i].cells[0].text = feat
    eng_table.rows[i].cells[1].text = dtype
    eng_table.rows[i].cells[2].text = desc

# ============================================================================
# 4. DATA GAP ANALYSIS - MISSING VALUES
# ============================================================================
doc.add_page_break()
doc.add_heading('4. Data Gap Analysis - Missing Values', level=1)

doc.add_paragraph(
    'A thorough understanding of missing data is essential for any data science project. Missing values can '
    'introduce bias, reduce statistical power, and limit the features that can be implemented in the final '
    'system. This section provides a comprehensive analysis of missing data patterns in the LinkedIn influencer dataset.'
)

doc.add_heading('4.1 Missing Values Summary', level=2)
doc.add_paragraph(
    'The following table presents a complete inventory of missing values across all 19 features in the dataset, '
    'sorted by the percentage of missing data from highest to lowest.'
)

# Create comprehensive missing values table
missing_table = doc.add_table(rows=12, cols=4)
missing_table.style = 'Table Grid'
missing_table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Column', 'Missing Count', 'Missing %', 'Data Type']
for i, header in enumerate(headers):
    missing_table.rows[0].cells[i].text = header
    missing_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

missing_data = [
    ('views', '34,012', '100.00%', 'float64'),
    ('votes', '33,926', '99.75%', 'object'),
    ('connections', '8,299', '24.40%', 'object'),
    ('media_type', '7,233', '21.27%', 'object'),
    ('location', '2,272', '6.68%', 'object'),
    ('content', '2,016', '5.93%', 'object'),
    ('followers', '42', '0.12%', 'float64'),
    ('time_spent', '1', '0.00%', 'object'),
    ('All other columns', '0', '0.00%', 'various'),
    ('(reactions, comments,', '', '', ''),
    ('hashtags, etc.)', '', '', ''),
]

for i, row_data in enumerate(missing_data, 1):
    if i < len(missing_table.rows):
        for j, val in enumerate(row_data):
            missing_table.rows[i].cells[j].text = val

doc.add_paragraph()

doc.add_heading('4.2 Visual Analysis of Missing Data', level=2)
doc.add_paragraph(
    'The visualization below provides two complementary views of the missing data pattern. The left panel '
    'shows a bar chart of missing percentages by column, making it easy to identify which features have '
    'the most significant data gaps. The right panel displays a heatmap of the missing data pattern across '
    'a sample of 1,000 records, revealing whether missing values are randomly distributed or follow specific patterns.'
)

add_image_with_caption(
    doc,
    os.path.join(BASE_PATH, 'missing_data_analysis.png'),
    'Missing Data Analysis - Bar Chart and Pattern Heatmap'
)

doc.add_heading('4.3 Interpretation of Missing Data Patterns', level=2)

doc.add_heading('Critical Missing Data (>50%)', level=3)
doc.add_paragraph(
    'Two features have critically high levels of missing data that effectively render them unusable for analysis:'
)

critical_missing = [
    ('views (100% missing)', 'The views column is entirely empty across all 34,012 records. This represents '
     'a significant data collection gap, as views/impressions are a key metric for understanding post reach '
     'and virality. Without this data, we cannot optimize content for maximum visibility or calculate '
     'view-to-engagement conversion rates. This gap likely resulted from LinkedIn API limitations or '
     'scraping restrictions that prevented access to view counts.'),
    ('votes (99.75% missing)', 'The votes column is nearly empty, with only 86 non-null values out of '
     '34,012 records. However, this is expected behavior rather than a data quality issue - votes only '
     'apply to poll-type posts, which represent only 0.25% of the dataset. The 86 valid entries correspond '
     'exactly to the 86 poll posts in the data.'),
]

for title_text, explanation in critical_missing:
    para = doc.add_paragraph()
    para.add_run(f'{title_text}: ').bold = True
    para.add_run(explanation)

doc.add_heading('Moderate Missing Data (10-50%)', level=3)
doc.add_paragraph(
    'Two features fall into the moderate missing data category, requiring careful consideration for analysis:'
)

moderate_missing = [
    ('connections (24.40% missing)', 'Approximately one quarter of records lack connection count data. '
     'However, analysis of available data reveals that this field has limited analytical value - 99.7% '
     'of non-null values show "500+" (the LinkedIn display cap), while only 80 records show the actual '
     'count of 171. This categorical nature limits its usefulness for quantitative analysis.'),
    ('media_type (21.27% missing)', 'About 21% of posts lack media type classification. This could indicate '
     'text-only posts without attachments, or could represent a data collection gap. Further investigation '
     'suggests these are likely text-only posts, as they still contain content but no media_url entries.'),
]

for title_text, explanation in moderate_missing:
    para = doc.add_paragraph()
    para.add_run(f'{title_text}: ').bold = True
    para.add_run(explanation)

doc.add_heading('Minor Missing Data (<10%)', level=3)
doc.add_paragraph(
    'Several features have minor missing data that can be handled through standard imputation or exclusion techniques:'
)

minor_missing = [
    ('location (6.68% missing)', 'Geographic location is missing for 2,272 records. This prevents geographic '
     'analysis for a subset of data but does not significantly impact content analysis features.'),
    ('content (5.93% missing)', 'Post content text is missing for 2,016 records. While this is a core feature, '
     'the 94% availability rate is sufficient for training NLP models. Missing content may indicate posts '
     'that were primarily media-based (images/videos) with minimal or no text.'),
    ('followers (0.12% missing)', 'Only 42 records lack follower count data, representing negligible missing '
     'data that can be easily imputed using the influencer\'s other posts.'),
]

for title_text, explanation in minor_missing:
    para = doc.add_paragraph()
    para.add_run(f'{title_text}: ').bold = True
    para.add_run(explanation)

# ============================================================================
# 5. DATA QUALITY ASSESSMENT
# ============================================================================
doc.add_page_break()
doc.add_heading('5. Data Quality Assessment', level=1)

doc.add_paragraph(
    'Beyond missing values, data quality encompasses issues such as duplicate records, inconsistent formatting, '
    'outliers, and data type mismatches. This section examines these aspects to ensure the dataset is suitable '
    'for building reliable machine learning models.'
)

doc.add_heading('5.1 Duplicate Records Analysis', level=2)
doc.add_paragraph(
    'Duplicate records can artificially inflate dataset size and skew analysis results. We examined the dataset '
    'for both exact duplicate rows and duplicate content.'
)

dup_table = doc.add_table(rows=3, cols=3)
dup_table.style = 'Table Grid'
dup_headers = ['Duplicate Type', 'Count', 'Percentage']
for i, h in enumerate(dup_headers):
    dup_table.rows[0].cells[i].text = h
    dup_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

dup_data = [
    ('Exact duplicate rows', '0', '0.00%'),
    ('Duplicate content text', '757', '2.37%'),
]
for i, (dtype, count, pct) in enumerate(dup_data, 1):
    dup_table.rows[i].cells[0].text = dtype
    dup_table.rows[i].cells[1].text = count
    dup_table.rows[i].cells[2].text = pct

doc.add_paragraph()
doc.add_paragraph(
    'The dataset contains no exact duplicate rows, indicating good data collection practices. However, '
    '757 posts (2.37%) share identical content text. This could represent: (1) influencers reposting '
    'their own successful content, (2) common phrases or templates used across posts, or (3) cross-posted '
    'content. For model training, these duplicates should be considered - they may need to be deduplicated '
    'or weighted appropriately depending on the modeling approach.'
)

doc.add_heading('5.2 Numerical Column Statistics', level=2)
doc.add_paragraph(
    'Understanding the statistical properties of numerical features is essential for feature engineering '
    'and model development. The table below presents comprehensive descriptive statistics for key numerical columns.'
)

stats_table = doc.add_table(rows=8, cols=6)
stats_table.style = 'Table Grid'

stats_headers = ['Statistic', 'Followers', 'Num Hashtags', 'Reactions', 'Comments', 'Eng. Rate*']
for i, h in enumerate(stats_headers):
    stats_table.rows[0].cells[i].text = h
    stats_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

stats_data = [
    ('Count', '33,970', '34,012', '34,012', '34,012', '33,970'),
    ('Mean', '1,125,922', '2.10', '473', '27', '0.77'),
    ('Std Dev', '3,057,750', '3.52', '4,164', '216', '3.91'),
    ('Min', '171', '0', '0', '0', '0.00'),
    ('Median (50%)', '408,254', '0', '36', '2', '0.11'),
    ('Max', '18,289,351', '48', '391,498', '32,907', '-'),
    ('Skewness', 'High +', 'High +', 'High +', 'High +', 'High +'),
]

for i, row_data in enumerate(stats_data, 1):
    for j, val in enumerate(row_data):
        stats_table.rows[i].cells[j].text = val

doc.add_paragraph()
doc.add_paragraph('*Engagement Rate = Reactions per 1,000 followers')
doc.add_paragraph()

doc.add_heading('Key Observations from Statistics', level=3)

stats_observations = [
    ('Highly Skewed Distributions', 'All engagement metrics show extreme positive skewness, with means '
     'significantly higher than medians. For example, the mean reactions (473) is over 13 times the median (36). '
     'This indicates a small number of viral posts driving up averages, while most posts receive modest engagement.'),
    ('Wide Range of Followers', 'Influencer follower counts span from 171 to over 18 million, representing '
     'a 100,000x range. This diversity is valuable for understanding how engagement scales with audience size.'),
    ('Hashtag Usage Patterns', 'The median hashtag count is 0, meaning more than half of posts use no hashtags. '
     'The average of 2.1 hashtags is pulled up by a minority of posts with heavy hashtag usage (max 48).'),
    ('Zero-Value Prevalence', 'A significant portion of posts have zero comments (32%) or zero reactions (3.1%), '
     'indicating that even influencer content doesn\'t always generate engagement.'),
]

for title_text, observation in stats_observations:
    para = doc.add_paragraph()
    para.add_run(f'{title_text}: ').bold = True
    para.add_run(observation)

doc.add_heading('5.3 Connections Field Analysis', level=2)
doc.add_paragraph(
    'The connections field exhibits unusual characteristics that warrant special attention. LinkedIn displays '
    '"500+" for users who have reached the connection limit, rather than showing the actual count.'
)

conn_table = doc.add_table(rows=4, cols=3)
conn_table.style = 'Table Grid'
conn_headers = ['Value', 'Count', 'Percentage']
for i, h in enumerate(conn_headers):
    conn_table.rows[0].cells[i].text = h
    conn_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

conn_data = [
    ('500+', '25,633', '75.4%'),
    ('171', '80', '0.2%'),
    ('Missing', '8,299', '24.4%'),
]
for i, row in enumerate(conn_data, 1):
    for j, val in enumerate(row):
        conn_table.rows[i].cells[j].text = val

doc.add_paragraph()
doc.add_paragraph(
    'This analysis reveals that the connections field has limited analytical utility. The overwhelming majority '
    '(99.7%) of non-null values are capped at "500+", making it essentially a binary indicator rather than '
    'a continuous variable. For modeling purposes, this field should either be: (1) converted to a binary '
    'feature (has_500_plus_connections), (2) excluded from analysis, or (3) treated as categorical.'
)

doc.add_heading('5.4 Time Field Analysis', level=2)
doc.add_paragraph(
    'The time_spent field captures when posts were published, but in a relative format rather than absolute timestamps.'
)

time_table = doc.add_table(rows=11, cols=2)
time_table.style = 'Table Grid'
time_headers = ['Time Value', 'Count']
for i, h in enumerate(time_headers):
    time_table.rows[0].cells[i].text = h
    time_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

time_data = [
    ('1 year ago', '7,753'),
    ('2 years ago', '5,728'),
    ('3 years ago', '3,759'),
    ('4 years ago', '2,126'),
    ('3 months ago', '1,456'),
    ('2 months ago', '1,448'),
    ('4 months ago', '1,279'),
    ('10 months ago', '1,247'),
    ('11 months ago', '1,133'),
    ('Other values', '~9,000'),
]
for i, row in enumerate(time_data, 1):
    for j, val in enumerate(row):
        time_table.rows[i].cells[j].text = val

doc.add_paragraph()
doc.add_paragraph(
    'The relative time format presents a significant limitation for temporal analysis. Without absolute timestamps, '
    'we cannot: (1) determine optimal posting times (day of week, hour of day), (2) analyze engagement trends over '
    'time, (3) account for LinkedIn algorithm changes, or (4) perform time-series forecasting. The data does show '
    'that posts span approximately 4+ years, providing good temporal diversity.'
)

# ============================================================================
# 6. DISTRIBUTION ANALYSIS
# ============================================================================
doc.add_page_break()
doc.add_heading('6. Distribution Analysis of Key Metrics', level=1)

doc.add_paragraph(
    'Understanding the distribution of engagement metrics is crucial for several reasons: (1) it informs '
    'the choice of statistical methods and machine learning algorithms, (2) it helps identify appropriate '
    'data transformations, and (3) it reveals the natural variation in post performance that models must capture.'
)

doc.add_heading('6.1 Engagement Metrics Distribution', level=2)
doc.add_paragraph(
    'The following visualization shows the distribution of four key metrics: reactions, comments, views, '
    'and followers. Due to the extreme skewness of these distributions, a logarithmic transformation '
    '(log(x+1)) was applied for visualization purposes.'
)

add_image_with_caption(
    doc,
    os.path.join(BASE_PATH, 'engagement_distributions.png'),
    'Distribution of Engagement Metrics (Log Scale)'
)

doc.add_heading('Interpretation of Engagement Distributions', level=3)

dist_interpretations = [
    ('Reactions Distribution (Top Left)', 'The log-transformed reactions show a roughly normal distribution '
     'centered around log(36) ≈ 3.6, indicating that most posts receive between 10-100 reactions. The long '
     'right tail represents viral posts with thousands to hundreds of thousands of reactions. This suggests '
     'a log-normal distribution is appropriate for modeling.'),
    ('Comments Distribution (Top Right)', 'Comments follow a similar log-normal pattern but are shifted left, '
     'reflecting that comments are less common than reactions. The peak around log(2) ≈ 0.7 indicates that '
     'most posts receive 0-5 comments. The 32% of posts with zero comments appear as a spike at the origin.'),
    ('Views Distribution (Bottom Left)', 'This panel is empty because views data is 100% missing. This '
     'represents a critical data gap that prevents reach optimization.'),
    ('Followers Distribution (Bottom Right)', 'Follower counts span a wide range on the log scale, from '
     'around 5 (≈150 followers) to 17 (≈18 million followers). The distribution appears bimodal, suggesting '
     'two distinct groups of influencers in the dataset.'),
]

for title_text, interpretation in dist_interpretations:
    para = doc.add_paragraph()
    para.add_run(f'{title_text}: ').bold = True
    para.add_run(interpretation)

doc.add_heading('6.2 Outlier Detection', level=2)
doc.add_paragraph(
    'Outliers can significantly impact model performance and must be carefully identified and handled. '
    'The box plots below show the distribution of each metric on a log scale, highlighting potential outliers '
    'as points beyond the whiskers.'
)

add_image_with_caption(
    doc,
    os.path.join(BASE_PATH, 'outlier_boxplots.png'),
    'Outlier Detection Box Plots (Log Scale)'
)

doc.add_paragraph(
    'The box plots reveal several important insights about outliers in the dataset:'
)

outlier_insights = [
    'All metrics show numerous outliers on the high end, represented by dots above the upper whiskers.',
    'These outliers represent viral posts that significantly outperform typical content.',
    'For reactions, the interquartile range (IQR) spans roughly log(7) to log(143), or 7-143 reactions.',
    'Extreme outliers include posts with 391,498 reactions (max) and 32,907 comments (max).',
    'For modeling, these outliers could be: (1) capped/winsorized, (2) log-transformed, or (3) modeled separately.',
]

for insight in outlier_insights:
    doc.add_paragraph(insight, style='List Bullet')

# ============================================================================
# 7. CONTENT ANALYSIS
# ============================================================================
doc.add_page_break()
doc.add_heading('7. Content Analysis', level=1)

doc.add_paragraph(
    'Content characteristics play a crucial role in determining post engagement. This section analyzes '
    'the textual properties of posts, including length, word count, hashtag usage, and media type distribution. '
    'These insights will inform the content optimization features of the final system.'
)

doc.add_heading('7.1 Content Length Statistics', level=2)
doc.add_paragraph(
    'Post length can impact readability, engagement, and algorithmic visibility. The following analysis '
    'examines both character count and word count distributions.'
)

length_table = doc.add_table(rows=5, cols=3)
length_table.style = 'Table Grid'
length_headers = ['Metric', 'Characters', 'Words']
for i, h in enumerate(length_headers):
    length_table.rows[0].cells[i].text = h
    length_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

length_data = [
    ('Average', '308', '49'),
    ('Median', '210', '33'),
    ('Maximum', '1,394', '260'),
    ('Empty Posts', '2,016 (5.9%)', '-'),
]
for i, row in enumerate(length_data, 1):
    for j, val in enumerate(row):
        length_table.rows[i].cells[j].text = val

doc.add_paragraph()

add_image_with_caption(
    doc,
    os.path.join(BASE_PATH, 'content_length_distribution.png'),
    'Distribution of Content Length (Characters and Words)'
)

doc.add_paragraph(
    'The content length distributions reveal several patterns relevant to content optimization:'
)

length_insights = [
    ('Moderate Length Preference', 'The median post contains 210 characters (~33 words), suggesting that '
     'successful influencers tend to keep posts concise. This aligns with LinkedIn best practices recommending '
     '150-300 characters for optimal engagement.'),
    ('Right-Skewed Distribution', 'Both character and word count distributions are right-skewed, with a long tail '
     'of longer posts. Some influencers write detailed, longer-form content up to 1,394 characters.'),
    ('Empty Content', '5.9% of posts have no text content. These are likely media-focused posts (images, videos) '
     'that rely on visual content rather than text to convey their message.'),
]

for title_text, insight in length_insights:
    para = doc.add_paragraph()
    para.add_run(f'{title_text}: ').bold = True
    para.add_run(insight)

doc.add_heading('7.2 Hashtag Analysis', level=2)
doc.add_paragraph(
    'Hashtags can increase post visibility by connecting content to broader conversations and making posts '
    'discoverable to users following specific topics. This analysis examines hashtag usage patterns among influencers.'
)

hashtag_stats = [
    ('Posts with hashtags', '14,405 (42.4%)'),
    ('Posts without hashtags', '19,607 (57.6%)'),
    ('Average hashtags per post', '2.10'),
    ('Maximum hashtags in a post', '48'),
    ('Most common count', '0 (no hashtags)'),
]

for stat, value in hashtag_stats:
    para = doc.add_paragraph(style='List Bullet')
    para.add_run(f'{stat}: ').bold = True
    para.add_run(value)

doc.add_paragraph()

add_image_with_caption(
    doc,
    os.path.join(BASE_PATH, 'hashtag_distribution.png'),
    'Distribution of Hashtag Count per Post'
)

doc.add_paragraph(
    'The hashtag distribution reveals an interesting pattern: the majority of influencer posts (57.6%) use no '
    'hashtags at all. Among posts that do use hashtags, 1-5 hashtags is most common. This challenges the common '
    'advice to always include hashtags - successful influencers may rely more on their established audience and '
    'content quality than hashtag-based discovery. However, the correlation between hashtag count and engagement '
    'should be examined before drawing conclusions.'
)

doc.add_heading('7.3 Media Type Distribution', level=2)
doc.add_paragraph(
    'The type of media attached to a post can significantly impact its performance. LinkedIn supports various '
    'media types including articles, images, videos, documents, and polls.'
)

add_image_with_caption(
    doc,
    os.path.join(BASE_PATH, 'media_type_distribution.png'),
    'Distribution of Media Types in Posts'
)

doc.add_paragraph(
    'The pie chart reveals the following media type distribution:'
)

media_breakdown = [
    ('Articles (44.5%)', 'Nearly half of all posts include article links, reflecting LinkedIn\'s positioning '
     'as a platform for sharing professional insights and external content.'),
    ('Images (25.6%)', 'About a quarter of posts feature images, used for infographics, quotes, and visual storytelling.'),
    ('No Media (21.3%)', 'Over one-fifth of posts are text-only, demonstrating that compelling text content can '
     'stand on its own without media attachments.'),
    ('Videos (7.9%)', 'Video content represents a smaller but significant portion of posts. As we\'ll see in the '
     'engagement analysis, videos often achieve the highest engagement despite their lower frequency.'),
    ('Other (0.7%)', 'Documents, polls, newsletters, and other media types make up a small fraction of posts.'),
]

for title_text, description in media_breakdown:
    para = doc.add_paragraph()
    para.add_run(f'{title_text}: ').bold = True
    para.add_run(description)

# ============================================================================
# 8. ENGAGEMENT ANALYSIS
# ============================================================================
doc.add_page_break()
doc.add_heading('8. Engagement Analysis', level=1)

doc.add_paragraph(
    'Engagement metrics are the primary target variables for the TrendPilot system. Understanding how '
    'engagement varies across different content types and influencers is essential for building effective '
    'recommendation models. This section provides detailed analysis of engagement patterns.'
)

doc.add_heading('8.1 Engagement by Media Type', level=2)
doc.add_paragraph(
    'Different media types generate significantly different levels of engagement. The following analysis '
    'compares average and median engagement metrics across all media types in the dataset.'
)

eng_media_table = doc.add_table(rows=9, cols=6)
eng_media_table.style = 'Table Grid'

eng_headers = ['Media Type', 'Count', 'Avg Reactions', 'Med. Reactions', 'Avg Comments', 'Med. Comments']
for i, h in enumerate(eng_headers):
    eng_media_table.rows[0].cells[i].text = h
    eng_media_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

eng_data = [
    ('Video', '2,690', '866', '114', '70', '12.5'),
    ('Image', '8,708', '824', '87.5', '40', '6'),
    ('Article', '15,144', '161', '13', '10', '1'),
    ('Document', '113', '95', '51', '29', '9'),
    ('Poll', '86', '58', '29.5', '34', '22.5'),
    ('Entity', '32', '50', '13.5', '4', '0'),
    ('Newsletter', '4', '17.5', '19', '1', '1'),
    ('View', '2', '4', '4', '0', '0'),
]

for i, row in enumerate(eng_data, 1):
    for j, val in enumerate(row):
        eng_media_table.rows[i].cells[j].text = val

doc.add_paragraph()

add_image_with_caption(
    doc,
    os.path.join(BASE_PATH, 'engagement_by_media.png'),
    'Engagement Comparison by Media Type (Log Scale)'
)

doc.add_heading('Key Insights from Media Type Analysis', level=3)

media_insights = [
    ('Video Dominates Engagement', 'Video content generates the highest average engagement with 866 reactions '
     'and 70 comments per post. The median video post receives 114 reactions - nearly 9x the median article. '
     'This strongly suggests that users should prioritize video content when seeking maximum engagement.'),
    ('Images Outperform Articles', 'Image posts average 824 reactions compared to 161 for articles. Despite '
     'articles being the most common media type (44.5% of posts), they generate the lowest engagement per post '
     'among the major categories. This represents a potential optimization opportunity.'),
    ('Polls Drive Comments', 'While polls have lower reaction counts, they generate disproportionately high '
     'comment engagement (median 22.5 comments). Polls may be effective for starting conversations and '
     'increasing audience interaction.'),
    ('Document Performance', 'Documents (PDFs, slideshows) show solid performance with median 51 reactions, '
     'suggesting they\'re effective for sharing detailed professional content.'),
]

for title_text, insight in media_insights:
    para = doc.add_paragraph()
    para.add_run(f'{title_text}: ').bold = True
    para.add_run(insight)

doc.add_heading('8.2 Engagement Rate Analysis', level=2)
doc.add_paragraph(
    'Raw engagement counts can be misleading because influencers with larger followings naturally receive '
    'more engagement. Engagement rate normalizes for audience size, providing a fairer comparison of content '
    'effectiveness. We calculate engagement rate as reactions per 1,000 followers.'
)

rate_table = doc.add_table(rows=4, cols=2)
rate_table.style = 'Table Grid'
rate_headers = ['Metric', 'Value']
for i, h in enumerate(rate_headers):
    rate_table.rows[0].cells[i].text = h
    rate_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

rate_data = [
    ('Mean Engagement Rate', '0.77 reactions per 1K followers'),
    ('Median Engagement Rate', '0.11 reactions per 1K followers'),
    ('Standard Deviation', '3.91'),
]
for i, row in enumerate(rate_data, 1):
    for j, val in enumerate(row):
        rate_table.rows[i].cells[j].text = val

doc.add_paragraph()
doc.add_paragraph(
    'The large gap between mean (0.77) and median (0.11) engagement rates indicates high variability in content '
    'performance. Even among established influencers, most posts achieve modest engagement rates, while a small '
    'percentage of posts dramatically outperform the rest. This suggests that content quality and relevance '
    'matter more than follower count alone.'
)

# ============================================================================
# 9. CORRELATION ANALYSIS
# ============================================================================
doc.add_page_break()
doc.add_heading('9. Correlation Analysis', level=1)

doc.add_paragraph(
    'Understanding the relationships between different features is essential for feature engineering and '
    'identifying predictors of engagement. This section examines correlations between key metrics using '
    'Pearson correlation coefficients.'
)

doc.add_heading('9.1 Correlation Matrix', level=2)

add_image_with_caption(
    doc,
    os.path.join(BASE_PATH, 'correlation_matrix.png'),
    'Correlation Matrix of Key Metrics'
)

doc.add_heading('9.2 Key Correlation Findings', level=2)

doc.add_heading('Strong Correlations', level=3)
doc.add_paragraph(
    'Reactions and Comments (r = 0.823): ').add_run(
    'The strongest correlation in the dataset exists between reactions and comments. This makes intuitive sense - '
    'posts that attract reactions also tend to generate discussion. This high correlation suggests that these two '
    'metrics could potentially be combined into a single "engagement score" for modeling purposes, reducing '
    'dimensionality while preserving signal.').bold = False

doc.add_heading('Moderate Correlations', level=3)
moderate_corrs = [
    ('Followers vs Reactions (r = 0.242)', 'A weak positive correlation exists between follower count and reactions. '
     'Larger audiences naturally lead to more reactions, but the moderate strength suggests content quality matters '
     'more than audience size.'),
    ('Followers vs Comments (r = 0.195)', 'Similar to reactions, comments show a weak positive correlation with '
     'followers. The slightly lower coefficient suggests that comments depend more on content that prompts discussion '
     'rather than raw audience size.'),
]

for title_text, description in moderate_corrs:
    para = doc.add_paragraph()
    para.add_run(f'{title_text}: ').bold = True
    para.add_run(description)

doc.add_heading('Negligible Correlations', level=3)
doc.add_paragraph(
    'Several features show negligible correlation with engagement, which is itself an important finding:'
)

negligible_corrs = [
    ('Num Hashtags vs Reactions (r = -0.042)', 'Contrary to popular belief, hashtag count shows almost no '
     'correlation with reactions. In fact, the very slight negative correlation suggests that excessive hashtag '
     'use might even slightly hurt engagement.'),
    ('Content Length vs Reactions (r = -0.027)', 'Post length has virtually no impact on engagement, suggesting '
     'that both short and long posts can be equally successful if the content is compelling.'),
    ('Word Count vs Comments (r = 0.025)', 'Similarly, word count shows no meaningful correlation with comments, '
     'indicating that prompting discussion depends on content substance rather than length.'),
]

for title_text, description in negligible_corrs:
    para = doc.add_paragraph()
    para.add_run(f'{title_text}: ').bold = True
    para.add_run(description)

doc.add_heading('9.3 Implications for Modeling', level=2)
doc.add_paragraph(
    'The correlation analysis has several important implications for building engagement prediction models:'
)

modeling_implications = [
    'Reactions and comments can potentially be combined into a single target variable due to their strong correlation.',
    'Follower count should be included as a feature but may benefit from interaction terms or normalization.',
    'Hashtag count and content length are unlikely to be strong predictors on their own.',
    'Media type (categorical) and content semantics (from NLP) may be more important predictors than simple numerical features.',
    'Feature engineering should focus on content quality metrics rather than quantity metrics.',
]

for impl in modeling_implications:
    doc.add_paragraph(impl, style='List Bullet')

# ============================================================================
# 10. TOP PERFORMERS ANALYSIS
# ============================================================================
doc.add_page_break()
doc.add_heading('10. Top Performers Analysis', level=1)

doc.add_paragraph(
    'Analyzing the top-performing influencers provides insights into what success looks like on LinkedIn and '
    'helps validate that the data captures meaningful variation in performance. This section examines the '
    'highest-engagement influencers and their posting patterns.'
)

doc.add_heading('10.1 Top Influencers by Average Engagement', level=2)
doc.add_paragraph(
    'The following table shows the top 10 influencers ranked by average reactions per post, filtered to include '
    'only those with at least 5 posts to ensure statistical reliability.'
)

top_table = doc.add_table(rows=11, cols=6)
top_table.style = 'Table Grid'

top_headers = ['Rank', 'Name', 'Avg Reactions', 'Total Posts', 'Followers', 'Avg Comments']
for i, h in enumerate(top_headers):
    top_table.rows[0].cells[i].text = h
    top_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

top_data = [
    ('1', 'Simon Sinek', '16,641', '264', '4.2M', '441'),
    ('2', 'Richard Branson', '5,186', '975', '18.3M', '236'),
    ('3', 'Kevin O\'Leary', '2,784', '424', '2.8M', '193'),
    ('4', 'Ian Bremmer', '2,158', '590', '3.7M', '99'),
    ('5', 'Vani Kola', '1,166', '949', '1.2M', '46'),
    ('6', 'Quentin M. Allums', '490', '366', '66K', '119'),
    ('7', 'Tom Goodwin', '375', '1,272', '719K', '67'),
    ('8', 'James Altucher', '338', '1,125', '1.3M', '27'),
    ('9', 'Natalie Riso', '248', '195', '406K', '22'),
    ('10', 'Tai T.', '242', '251', '348K', '21'),
]

for i, row in enumerate(top_data, 1):
    for j, val in enumerate(row):
        top_table.rows[i].cells[j].text = val

doc.add_paragraph()

doc.add_heading('10.2 Observations from Top Performers', level=2)

top_observations = [
    ('Exceptional Performance', 'Simon Sinek stands out dramatically with an average of 16,641 reactions per post - '
     'over 3x the second-place Richard Branson. His content consistently resonates with his audience, making his '
     'posts valuable case studies for content optimization.'),
    ('Follower Count Not Deterministic', 'Richard Branson has the most followers (18.3M) but ranks second in '
     'engagement. Meanwhile, Quentin M. Allums achieves high engagement (490 avg reactions) with only 66K followers, '
     'demonstrating that audience engagement depends on more than audience size.'),
    ('Posting Frequency Varies', 'Top performers show varied posting frequencies from 195 posts (Natalie Riso) '
     'to 1,272 posts (Tom Goodwin). There\'s no clear correlation between posting volume and average engagement, '
     'suggesting quality matters more than quantity.'),
    ('Comments vs Reactions Ratio', 'Simon Sinek also leads in comments with 441 average comments per post, but '
     'the reactions-to-comments ratio varies across influencers. Some generate more discussion relative to their '
     'reactions, indicating different content styles.'),
]

for title_text, observation in top_observations:
    para = doc.add_paragraph()
    para.add_run(f'{title_text}: ').bold = True
    para.add_run(observation)

# ============================================================================
# 11. DATA GAP RISK ASSESSMENT
# ============================================================================
doc.add_page_break()
doc.add_heading('11. Data Gap Risk Assessment', level=1)

doc.add_paragraph(
    'This section provides a comprehensive risk assessment of identified data gaps, categorizing each by '
    'severity and analyzing the specific impact on planned system features. Understanding these risks is '
    'essential for project planning and setting appropriate expectations.'
)

doc.add_heading('11.1 Risk Assessment Matrix', level=2)

risk_table = doc.add_table(rows=9, cols=4)
risk_table.style = 'Table Grid'

risk_headers = ['Data Gap', 'Missing %', 'Severity', 'Impact on Project']
for i, h in enumerate(risk_headers):
    risk_table.rows[0].cells[i].text = h
    risk_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

risk_data = [
    ('Views data', '100%', 'HIGH', 'Cannot optimize for reach/virality'),
    ('Absolute timestamps', '100%*', 'HIGH', 'Cannot recommend posting times'),
    ('Hashtag followers', '100%**', 'MEDIUM', 'Cannot assess hashtag potential'),
    ('Connections', '24.4%', 'LOW', 'Field has limited utility anyway'),
    ('Media type', '21.3%', 'MEDIUM', 'Reduced sample for media analysis'),
    ('Location', '6.7%', 'MEDIUM', 'Limited geographic targeting'),
    ('Content', '5.9%', 'LOW', 'Minimal impact - 94% available'),
    ('Followers', '0.12%', 'LOW', 'Easily imputable'),
]

for i, row in enumerate(risk_data, 1):
    for j, val in enumerate(row):
        risk_table.rows[i].cells[j].text = val

doc.add_paragraph()
doc.add_paragraph('*Timestamps exist but in relative format only (e.g., "2 days ago")')
doc.add_paragraph('**Column exists but all values are 0')

doc.add_heading('11.2 Detailed Risk Analysis', level=2)

doc.add_heading('HIGH Severity Risks', level=3)

high_risks = [
    ('Views Data Gap',
     'The complete absence of views data represents the most significant data gap. Views/impressions are a '
     'fundamental metric for understanding post reach and calculating conversion rates (views-to-engagement). '
     'Without this data, we cannot: (1) distinguish between posts that reached many people but received low '
     'engagement vs posts with limited reach, (2) optimize content for visibility in the LinkedIn feed, '
     '(3) calculate impression-based engagement rates for fairer comparisons.',
     'MITIGATION: Focus on engagement optimization rather than reach optimization. Use follower count as a '
     'proxy for potential reach. Consider this a limitation in the final system documentation.'),
    ('Timestamp Format Gap',
     'While time data exists, it\'s in relative format ("1 week ago") rather than absolute timestamps. This '
     'prevents temporal analysis including: (1) optimal posting day of week, (2) optimal posting hour, '
     '(3) engagement decay patterns over time, (4) seasonal or trending topic analysis.',
     'MITIGATION: Acknowledge this limitation. For future data collection, prioritize capturing absolute '
     'timestamps. Consider using LinkedIn\'s API for time-sensitive analysis.'),
]

for title_text, description, mitigation in high_risks:
    para = doc.add_paragraph()
    para.add_run(f'{title_text}\n').bold = True
    para.add_run(f'{description}\n\n')
    para.add_run(f'{mitigation}').italic = True
    doc.add_paragraph()

doc.add_heading('MEDIUM Severity Risks', level=3)

medium_risks = [
    ('Media Type Gap (21.3%)',
     'About one-fifth of posts lack media type classification. This reduces the sample size for media-based '
     'analysis and may introduce bias if certain media types are more likely to have missing values.',
     'MITIGATION: Analyze patterns in missing media types. Consider treating "No Media" as its own category. '
     'Ensure models handle missing values appropriately.'),
    ('Location Gap (6.7%)',
     'Geographic data is missing for some records, limiting geographic targeting features.',
     'MITIGATION: Geographic features may not be critical for content optimization. Consider deprioritizing '
     'location-based recommendations.'),
]

for title_text, description, mitigation in medium_risks:
    para = doc.add_paragraph()
    para.add_run(f'{title_text}\n').bold = True
    para.add_run(f'{description}\n\n')
    para.add_run(f'{mitigation}').italic = True
    doc.add_paragraph()

# ============================================================================
# 12. IMPACT ON FINAL PROJECT
# ============================================================================
doc.add_page_break()
doc.add_heading('12. Impact Analysis on Final Project', level=1)

doc.add_paragraph(
    'The TrendPilot system aims to help users create engaging LinkedIn posts. This section evaluates the '
    'feasibility and scope of each planned feature based on the data availability and quality findings '
    'from this EDA.'
)

doc.add_heading('12.1 Feature Feasibility Assessment', level=2)

# Feature 1
doc.add_heading('Content Optimization Features', level=3)
feat1_table = doc.add_table(rows=4, cols=2)
feat1_table.style = 'Table Grid'
feat1_data = [
    ('Data Available', '94.1% of posts have content'),
    ('Status', 'FULLY FEASIBLE'),
    ('Confidence', 'HIGH'),
]
for i, (label, value) in enumerate(feat1_data):
    feat1_table.rows[i].cells[0].text = label
    feat1_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    feat1_table.rows[i].cells[1].text = value

doc.add_paragraph()
doc.add_paragraph('Capabilities enabled by available data:')
content_capabilities = [
    'Analyze text patterns, structure, and writing style of high-engagement posts',
    'Build NLP models to suggest content improvements',
    'Identify keywords, phrases, and topics associated with high engagement',
    'Recommend optimal content length based on engagement correlations',
    'Detect sentiment and tone patterns in successful posts',
]
for cap in content_capabilities:
    doc.add_paragraph(cap, style='List Bullet')

# Feature 2
doc.add_heading('Hashtag Recommendation Features', level=3)
feat2_table = doc.add_table(rows=4, cols=2)
feat2_table.style = 'Table Grid'
feat2_data = [
    ('Data Available', '100% hashtag lists, 0% follower counts'),
    ('Status', 'PARTIALLY FEASIBLE'),
    ('Confidence', 'MEDIUM'),
]
for i, (label, value) in enumerate(feat2_data):
    feat2_table.rows[i].cells[0].text = label
    feat2_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    feat2_table.rows[i].cells[1].text = value

doc.add_paragraph()
doc.add_paragraph('Available capabilities:')
hashtag_available = [
    'Recommend hashtags based on content topic matching',
    'Analyze which hashtags correlate with higher engagement',
    'Suggest optimal number of hashtags based on data patterns',
]
for cap in hashtag_available:
    doc.add_paragraph(cap, style='List Bullet')

doc.add_paragraph('Unavailable due to data gaps:')
hashtag_unavailable = [
    'Recommend hashtags by reach potential (no follower data)',
    'Predict hashtag-driven visibility improvements',
]
for cap in hashtag_unavailable:
    doc.add_paragraph(cap, style='List Bullet')

# Feature 3
doc.add_heading('Media Type Recommendations', level=3)
feat3_table = doc.add_table(rows=4, cols=2)
feat3_table.style = 'Table Grid'
feat3_data = [
    ('Data Available', '78.7% of posts have media type'),
    ('Status', 'FULLY FEASIBLE'),
    ('Confidence', 'HIGH'),
]
for i, (label, value) in enumerate(feat3_data):
    feat3_table.rows[i].cells[0].text = label
    feat3_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    feat3_table.rows[i].cells[1].text = value

doc.add_paragraph()
doc.add_paragraph('Capabilities enabled by available data:')
media_capabilities = [
    'Recommend optimal media type based on content and goals',
    'Quantify engagement differences between media types',
    'Suggest video/image when higher engagement is the goal',
    'Identify when text-only posts are appropriate',
]
for cap in media_capabilities:
    doc.add_paragraph(cap, style='List Bullet')

# Feature 4
doc.add_heading('Engagement Prediction Model', level=3)
feat4_table = doc.add_table(rows=4, cols=2)
feat4_table.style = 'Table Grid'
feat4_data = [
    ('ML-Ready Rows', '26,107 (76.8%)'),
    ('Status', 'FULLY FEASIBLE'),
    ('Confidence', 'HIGH'),
]
for i, (label, value) in enumerate(feat4_data):
    feat4_table.rows[i].cells[0].text = label
    feat4_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    feat4_table.rows[i].cells[1].text = value

doc.add_paragraph()
doc.add_paragraph('Model capabilities:')
model_capabilities = [
    'Predict expected reactions based on content features',
    'Predict expected comments based on content features',
    'Provide confidence intervals for predictions',
    'Identify factors most predictive of engagement',
]
for cap in model_capabilities:
    doc.add_paragraph(cap, style='List Bullet')

doc.add_paragraph('Critical limitation:')
doc.add_paragraph('Cannot predict or optimize for views/reach due to 100% missing views data', style='List Bullet')

# Feature 5
doc.add_heading('Posting Time Recommendations', level=3)
feat5_table = doc.add_table(rows=4, cols=2)
feat5_table.style = 'Table Grid'
feat5_data = [
    ('Data Available', 'Relative time only'),
    ('Status', 'NOT FEASIBLE'),
    ('Confidence', 'N/A'),
]
for i, (label, value) in enumerate(feat5_data):
    feat5_table.rows[i].cells[0].text = label
    feat5_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    feat5_table.rows[i].cells[1].text = value

doc.add_paragraph()
doc.add_paragraph(
    'This feature cannot be implemented with the current data. The relative timestamp format ("2 weeks ago") '
    'does not provide the day-of-week or hour-of-day information needed for posting time optimization. '
    'This feature should be deprioritized or planned for a future phase with new data collection.'
)

# ============================================================================
# 13. RECOMMENDATIONS
# ============================================================================
doc.add_page_break()
doc.add_heading('13. Recommendations', level=1)

doc.add_paragraph(
    'Based on the findings of this EDA, the following recommendations are provided for the development '
    'of the TrendPilot system and for future data collection efforts.'
)

doc.add_heading('13.1 Feature Prioritization', level=2)

doc.add_heading('Tier 1: Focus Features (Sufficient Data)', level=3)
doc.add_paragraph('These features should be the primary focus of development:')
tier1_features = [
    ('Content Text Analysis', 'Build NLP models to analyze content patterns, suggest improvements, and predict engagement.'),
    ('Media Type Recommendations', 'Recommend optimal media types based on clear engagement differences in the data.'),
    ('Engagement Prediction', 'Develop ML models to predict reactions and comments based on content features.'),
    ('Hashtag Suggestions', 'Recommend relevant hashtags based on content, even without reach data.'),
]
for title_text, description in tier1_features:
    para = doc.add_paragraph()
    para.add_run(f'{title_text}: ').bold = True
    para.add_run(description)

doc.add_heading('Tier 2: Limited Features (Partial Data)', level=3)
doc.add_paragraph('These features can be implemented with acknowledged limitations:')
tier2_features = [
    ('Engagement Rate Normalization', 'Use follower count to normalize engagement, acknowledging some missing values.'),
    ('Content Length Guidance', 'Provide soft recommendations based on observed patterns, though correlation is weak.'),
]
for title_text, description in tier2_features:
    para = doc.add_paragraph()
    para.add_run(f'{title_text}: ').bold = True
    para.add_run(description)

doc.add_heading('Tier 3: Deprioritized Features (Insufficient Data)', level=3)
doc.add_paragraph('These features should be deprioritized or excluded:')
tier3_features = [
    ('Views/Reach Optimization', 'Cannot implement - views data is 100% missing.'),
    ('Posting Time Recommendations', 'Cannot implement - no absolute timestamps available.'),
    ('Geographic Targeting', 'Limited value - location data is sparse and may not impact content strategy.'),
    ('Hashtag Reach Prediction', 'Cannot implement - hashtag follower data is all zeros.'),
]
for title_text, description in tier3_features:
    para = doc.add_paragraph()
    para.add_run(f'{title_text}: ').bold = True
    para.add_run(description)

doc.add_heading('13.2 Data Collection Recommendations', level=2)
doc.add_paragraph('For future data collection phases, prioritize the following:')

data_recs = [
    ('Absolute Timestamps', 'Capture the exact date and time of each post to enable temporal analysis and posting time recommendations.'),
    ('Views/Impressions', 'Ensure views data is captured consistently. This may require different collection methods or API access.'),
    ('Hashtag Metadata', 'Collect hashtag follower counts when available to enable reach-based recommendations.'),
    ('Engagement Over Time', 'Track how engagement accumulates over time (1 hour, 24 hours, 1 week) to understand decay patterns.'),
    ('Profile Updates', 'Capture follower growth over time to understand audience building patterns.'),
]

for title_text, description in data_recs:
    para = doc.add_paragraph()
    para.add_run(f'{title_text}: ').bold = True
    para.add_run(description)

doc.add_heading('13.3 Model Development Recommendations', level=2)

model_recs = [
    ('Primary Target Variable', 'Use reactions + comments as the primary engagement metric. Given their strong correlation (r=0.823), consider combining them into a single score or predicting them jointly.'),
    ('Handle Skewness', 'Apply log transformation to engagement metrics before modeling. The extreme skewness of raw values will otherwise dominate model training.'),
    ('Engagement Rate', 'Consider using engagement rate (reactions/followers) as an alternative target to account for audience size differences.'),
    ('Feature Engineering', 'Focus on content-derived features (sentiment, topics, readability) rather than simple counts. Correlation analysis shows counts have limited predictive power.'),
    ('Missing Value Strategy', 'For media_type, treat missing as a "No Media" category. For other features, use appropriate imputation or exclusion depending on the modeling approach.'),
    ('Validation Strategy', 'Use influencer-stratified cross-validation to ensure models generalize across different posting styles.'),
]

for title_text, description in model_recs:
    para = doc.add_paragraph()
    para.add_run(f'{title_text}: ').bold = True
    para.add_run(description)

# ============================================================================
# 14. CONCLUSION
# ============================================================================
doc.add_page_break()
doc.add_heading('14. Conclusion', level=1)

doc.add_heading('14.1 Summary of Findings', level=2)

doc.add_paragraph(
    'This exploratory data analysis has provided a comprehensive assessment of the LinkedIn influencer dataset '
    'for the TrendPilot project. The dataset comprises 34,012 posts from 69 influential LinkedIn users, '
    'offering a substantial foundation for building engagement prediction and content optimization models.'
)

doc.add_paragraph(
    'The analysis revealed that the dataset achieves strong data quality for core features, with 95.4% '
    'completeness for the features most critical to engagement prediction (content, reactions, comments, '
    'media type, hashtags, and followers). This level of completeness is sufficient for developing reliable '
    'machine learning models.'
)

doc.add_paragraph(
    'However, significant gaps exist that limit certain capabilities. Most notably, views data is completely '
    'missing (100%), preventing reach optimization. Additionally, timestamps are stored in relative format, '
    'preventing temporal analysis and posting time recommendations. These limitations should be clearly '
    'communicated in the final system and addressed in future data collection efforts.'
)

doc.add_heading('14.2 Data Quality Score', level=2)

final_table = doc.add_table(rows=5, cols=2)
final_table.style = 'Table Grid'
final_data = [
    ('Overall Data Completeness', '86.7%'),
    ('Core Features Completeness', '95.4%'),
    ('Quality Assessment', 'GOOD'),
    ('Suitability for ML', 'SUITABLE'),
]
for i, (metric, value) in enumerate(final_data):
    final_table.rows[i].cells[0].text = metric
    final_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    final_table.rows[i].cells[1].text = value

doc.add_paragraph()

doc.add_heading('14.3 Key Takeaways', level=2)

takeaways = [
    'The dataset is suitable for building a LinkedIn post optimization system focused on content analysis, '
    'media type recommendations, and engagement prediction.',
    'Video and image content significantly outperform articles in engagement, suggesting clear recommendations '
    'for users seeking maximum impact.',
    'Reactions and comments are strongly correlated (r=0.823), enabling their combination into a unified engagement metric.',
    'Hashtag count and content length show negligible correlation with engagement, challenging conventional wisdom.',
    'Top performers like Simon Sinek achieve exceptional engagement through content quality, not just audience size.',
    'Future data collection should prioritize absolute timestamps and views data to enable currently infeasible features.',
]

for takeaway in takeaways:
    doc.add_paragraph(takeaway, style='List Bullet')

doc.add_heading('14.4 Next Steps', level=2)

next_steps = [
    'Proceed with feature engineering based on EDA insights',
    'Develop NLP pipelines for content analysis',
    'Build engagement prediction models using available features',
    'Create media type recommendation logic',
    'Document limitations for user transparency',
    'Plan future data collection to address identified gaps',
]

for i, step in enumerate(next_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_paragraph()
doc.add_paragraph(
    'This EDA provides a solid foundation for the development of TrendPilot. While some features will be '
    'limited by data gaps, the available data supports the core goal of helping users create more engaging '
    'LinkedIn content through data-driven recommendations.'
)

# Save document
output_path = '/Users/BirajMishra/work/playground/ms/capstone/capstone_trend_pilot/eda/LinkedIn_EDA_Report.docx'
doc.save(output_path)
print(f"Comprehensive report saved to: {output_path}")
